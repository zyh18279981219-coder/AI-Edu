from __future__ import annotations

import json
import math
import sys
import html
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from textwrap import wrap
from typing import Any

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
BACKEND = ROOT / "backend"
for item in (ROOT, BACKEND):
    value = str(item)
    if value not in sys.path:
        sys.path.insert(0, value)

from tools.docs.generate_live_design_artifacts import (  # noqa: E402
    MODULES,
    classify_operation,
    classify_table,
    load_database_schema,
    load_openapi_operations,
)


OUT_DIR = ROOT / "docs" / "detailed-design"
HIGH_SPEC_DIR = OUT_DIR / "high-spec"
FIG_DIR = HIGH_SPEC_DIR / "figures"
DRAWIO_DIR = HIGH_SPEC_DIR / "drawio"
DOCX_PATH = OUT_DIR / "AI-Education系统详细设计说明书-实现对齐高规格版-2026-06-29.docx"
MD_PATH = OUT_DIR / "AI-Education系统详细设计说明书-实现对齐高规格版-2026-06-29.md"
REQ_DOC = ROOT / "docs" / "requirements" / "AI-Education需求分析文档-图件美化版-2026-06-28.docx"
BASELINE_MD = OUT_DIR / "AI-Education系统详细设计说明书-开发基线版-2026-06-28.md"


PALETTE = {
    "blue": "#DDEBFF",
    "blue_border": "#2F66B3",
    "green": "#E2F3E8",
    "green_border": "#2F8A56",
    "orange": "#FFF0D8",
    "orange_border": "#C77814",
    "purple": "#ECE7FA",
    "purple_border": "#7256B5",
    "cyan": "#DDF4F7",
    "cyan_border": "#288C99",
    "red": "#FCE3E0",
    "red_border": "#C9443C",
    "gray": "#F4F6F8",
    "gray_border": "#7C8794",
    "ink": "#172033",
    "muted": "#52606D",
}


MODULE_GAPS = {
    "课程数字孪生与课程资源": [
        "课程底座发布态过滤需要继续代码级验收，下游不得读取未发布草稿节点、候选资源和候选能力映射。",
        "资源自动检索质量、RAG 入库和资源审核体验仍需按真实页面闭环。",
    ],
    "学生学习空间": [
        "需核对首页、课程资源、作业、测验、画像摘要和路径任务是否形成统一学生工作台。",
        "学生端不得暴露教师审核过程、证据权重调整和后台备注。",
    ],
    "在线测验": [
        "测验定义、发布和作答接口已存在，但题目版本、知识点绑定和异常测验排除规则需要补字段级说明。",
        "测验结果回流画像、诊断和教师看板的触发时机需要代码级验收。",
    ],
    "学生数字孪生": [
        "能力达成等级需要确认只读取教师确认并发布后的职业能力映射。",
        "证据下钻、快照追溯和学生/教师视图差异仍需前端验收。",
    ],
    "诊断智能体": [
        "诊断修正表当前为空，人工修正闭环需要补演示数据和页面入口。",
        "证据不足分支必须阻止强诊断进入正式路径或强干预。",
    ],
    "个性化学习路径": [
        "路径生成应由学生端触发和执行，教师端只进行干预；`username` 参数接口需做角色约束。",
        "路径版本的触发源、诊断依据、失效原因和回滚字段仍需补强。",
    ],
    "5E 教学智能体": [
        "5E 事件、交互统计和有效性记录当前为空，属于结构已建未产品化。",
        "5E 证据只能作为过程性辅助证据，不能替代测验和作业。",
    ],
    "作业与实践评测": [
        "作业覆盖知识点必须只在 `confirmed_by_teacher=1` 后影响叶子知识点画像。",
        "代码题稳定性、权限控制和状态流转需要继续专项验证。",
    ],
    "教师智能干预任务包": [
        "干预任务完成结果回流画像、路径和看板的闭环需要继续验收。",
        "`teacher/diagnose` 应调用统一诊断服务，避免形成第二套诊断口径。",
    ],
    "教师看板与教师数字孪生": [
        "所有看板接口必须按教师-学生或课程授权范围过滤。",
        "AI 建议只能教师手动触发，不得写成系统自动替教师决策。",
    ],
    "教学互动": [
        "公告、讨论、教研记录表当前多为空，需要补演示数据和前端验收。",
        "互动证据进入学生画像时需要保留来源、时间、课程节点和任务类型。",
    ],
    "行业情报与能力对接": [
        "3.12 产生岗位/能力候选，3.1 完成审核发布，设计中需避免职责混写。",
        "职业能力映射只允许使用教师确认发布后的叶子知识点关系。",
    ],
}


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = _font(38, True)
FONT_SUBTITLE = _font(26, True)
FONT_NODE = _font(24, False)
FONT_SMALL = _font(20, False)


def _draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str = PALETTE["ink"],
    max_chars: int = 12,
) -> None:
    lines: list[str] = []
    for part in text.split("\n"):
        lines.extend(wrap(part, width=max_chars) or [""])
    line_heights = []
    widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1] + 8)
    total_h = sum(line_heights)
    y = box[1] + (box[3] - box[1] - total_h) / 2
    for i, line in enumerate(lines):
        x = box[0] + (box[2] - box[0] - widths[i]) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_heights[i]


def _box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str,
    border: str,
    font: ImageFont.FreeTypeFont = FONT_NODE,
    radius: int = 18,
    max_chars: int = 12,
) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=border, width=3)
    _draw_centered_text(draw, xy, text, font, max_chars=max_chars)


def _arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#42526E",
    width: int = 4,
    label: str | None = None,
) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    pts = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 7), end[1] - size * math.sin(angle - math.pi / 7)),
        (end[0] - size * math.cos(angle + math.pi / 7), end[1] - size * math.sin(angle + math.pi / 7)),
    ]
    draw.polygon(pts, fill=color)
    if label:
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 26
        bbox = draw.textbbox((0, 0), label, font=FONT_SMALL)
        draw.rounded_rectangle((mx - 12, my - 8, mx + bbox[2] + 12, my + bbox[3] + 8), radius=8, fill="#FFFFFF", outline="#CBD5E1")
        draw.text((mx, my), label, font=FONT_SMALL, fill=PALETTE["muted"])


def _elbow_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    color: str = "#42526E",
    width: int = 4,
    label: str | None = None,
) -> None:
    draw.line(points, fill=color, width=width, joint="curve")
    _arrow(draw, points[-2], points[-1], color=color, width=width, label=label)


def _canvas(title: str, subtitle: str | None = None, size: tuple[int, int] = (2200, 1300)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, size[0], 110), fill="#F8FAFC")
    draw.text((70, 34), title, font=FONT_TITLE, fill=PALETTE["ink"])
    if subtitle:
        draw.text((70, 78), subtitle, font=FONT_SMALL, fill=PALETTE["muted"])
    return image, draw


def _save(image: Image.Image, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / f"{name}.png"
    image.save(path)
    return path


def fig_system_architecture() -> Path:
    image, draw = _canvas("图 1  AI-Education 运行架构与边界", "前端入口、FastAPI 服务、智能服务与 MySQL 数据底座分层")
    _box(draw, (80, 190, 600, 320), "学生端\n学习、画像、路径、互动", PALETTE["blue"], PALETTE["blue_border"], max_chars=16)
    _box(draw, (80, 390, 600, 520), "教师端\n课程建设、看板、干预", PALETTE["blue"], PALETTE["blue_border"], max_chars=16)
    _box(draw, (80, 590, 600, 720), "管理员端\n账号、课程、运行治理", PALETTE["blue"], PALETTE["blue_border"], max_chars=16)
    _box(draw, (760, 360, 1480, 500), "FastAPI 接口层\n164 个 OpenAPI operation", PALETTE["green"], PALETTE["green_border"], max_chars=22)
    _box(draw, (800, 610, 1090, 760), "业务服务\n课程/测验/作业/路径", PALETTE["green"], PALETTE["green_border"], max_chars=12)
    _box(draw, (1190, 610, 1480, 760), "智能服务\n诊断/5E/行业情报", PALETTE["purple"], PALETTE["purple_border"], max_chars=12)
    _box(draw, (800, 890, 1480, 1035), "MySQL ai_education_design\n42 张真实业务与支撑表", PALETTE["orange"], PALETTE["orange_border"], max_chars=24)
    _box(draw, (1640, 330, 2080, 470), "外部资源\nB站/YouTube/CSDN/文档", PALETTE["cyan"], PALETTE["cyan_border"], max_chars=18)
    _box(draw, (1640, 650, 2080, 790), "大模型与工具\nLLM/OCR/RAG/日志", PALETTE["cyan"], PALETTE["cyan_border"], max_chars=18)
    draw.line([(680, 255), (680, 655)], fill="#42526E", width=4)
    for y in (255, 455, 655):
        _arrow(draw, (600, y), (680, y))
    _arrow(draw, (680, 455), (760, 430))
    _elbow_arrow(draw, [(1120, 500), (1120, 560), (945, 560), (945, 610)])
    _elbow_arrow(draw, [(1120, 500), (1120, 560), (1335, 560), (1335, 610)])
    _arrow(draw, (945, 760), (1050, 890))
    _arrow(draw, (1335, 760), (1220, 890))
    _arrow(draw, (1480, 680), (1640, 400))
    _arrow(draw, (1480, 710), (1640, 720))
    return _save(image, "fig_1_system_architecture")


def fig_module_collaboration() -> Path:
    image, draw = _canvas("图 2  需求模块与实现模块协同关系", "按 3.1-3.12 职责边界组织，不按历史 URL 前缀机械归类")
    cols = [
        ("课程与行业底座", ["3.1 课程数字孪生", "3.12 行业情报", "资源与能力发布"]),
        ("学习证据采集", ["3.2 学习空间", "3.3 在线测验", "3.8 作业评测", "3.7 5E 辅助"]),
        ("画像诊断路径", ["3.4 学生数字孪生", "3.5 诊断智能体", "3.6 个性化路径"]),
        ("教师决策反馈", ["3.10 教师看板", "3.9 教师干预", "3.11 教学互动"]),
    ]
    xs = [80, 610, 1140, 1670]
    for idx, (title, nodes) in enumerate(cols):
        x = xs[idx]
        draw.rounded_rectangle((x, 180, x + 430, 1050), radius=24, fill="#F8FAFC", outline="#CBD5E1", width=3)
        draw.text((x + 35, 220), title, font=FONT_SUBTITLE, fill=PALETTE["ink"])
        for j, node in enumerate(nodes):
            y = 330 + j * 180
            fill, border = [(PALETTE["blue"], PALETTE["blue_border"]), (PALETTE["green"], PALETTE["green_border"]), (PALETTE["purple"], PALETTE["purple_border"]), (PALETTE["orange"], PALETTE["orange_border"])][idx]
            _box(draw, (x + 45, y, x + 385, y + 105), node, fill, border, max_chars=13)
    for i in range(3):
        _arrow(draw, (xs[i] + 430, 600), (xs[i + 1], 600), label=["已发布底座", "学习证据", "风险与任务"][i])
    _elbow_arrow(draw, [(1885, 880), (1885, 1120), (290, 1120), (290, 1050)], label="反馈优化")
    return _save(image, "fig_2_module_collaboration")


def fig_data_closed_loop() -> Path:
    image, draw = _canvas("图 3  课程底座—学习证据—画像诊断—教学反馈数据闭环", "强调发布、证据、诊断、路径和干预的闭环边界")
    nodes = [
        ((120, 230, 440, 350), "教师建设课程底座", PALETTE["blue"], PALETTE["blue_border"]),
        ((560, 230, 880, 350), "发布课程图谱\n资源与能力映射", PALETTE["green"], PALETTE["green_border"]),
        ((1000, 230, 1320, 350), "学生学习\n资源/测验/作业/互动", PALETTE["orange"], PALETTE["orange_border"]),
        ((1440, 230, 1760, 350), "学生数字孪生\n画像与快照", PALETTE["purple"], PALETTE["purple_border"]),
        ((1440, 520, 1760, 640), "诊断智能体\n薄弱点与原因", PALETTE["purple"], PALETTE["purple_border"]),
        ((1000, 520, 1320, 640), "学生个性化路径\n补学与执行", PALETTE["green"], PALETTE["green_border"]),
        ((560, 520, 880, 640), "教师看板\n班级风险与证据", PALETTE["blue"], PALETTE["blue_border"]),
        ((120, 520, 440, 640), "教师干预任务包\n审核后下发", PALETTE["blue"], PALETTE["blue_border"]),
        ((560, 820, 1320, 960), "结果回流：画像更新 / 路径调整 / 课程资源维护 / 教师画像事件", PALETTE["cyan"], PALETTE["cyan_border"]),
    ]
    for xy, text, fill, border in nodes:
        _box(draw, xy, text, fill, border, max_chars=16)
    for a, b in [((440, 290), (560, 290)), ((880, 290), (1000, 290)), ((1320, 290), (1440, 290)), ((1600, 350), (1600, 520)), ((1440, 580), (1320, 580)), ((1000, 580), (880, 580)), ((560, 580), (440, 580))]:
        _arrow(draw, a, b)
    _elbow_arrow(draw, [(280, 640), (280, 890), (560, 890)])
    _arrow(draw, (1320, 890), (1600, 640), label="证据更新")
    _elbow_arrow(draw, [(940, 820), (940, 750), (280, 750), (280, 640)], label="教学反馈")
    return _save(image, "fig_3_data_closed_loop")


def fig_database_domains(schema: dict[str, Any]) -> Path:
    image, draw = _canvas("图 4  MySQL 数据库分域与表结构现状", f"当前库 {schema['database']}：{len(schema['tables'])} 张表，行数为本地演示与验证快照")
    groups = [
        ("课程与资源", ["courses", "course_nodes", "resources", "course_ability_mappings"], PALETTE["blue"], PALETTE["blue_border"]),
        ("学习证据", ["quiz_attempts", "homework_submissions", "resource_learning_events", "user_interaction"], PALETTE["orange"], PALETTE["orange_border"]),
        ("画像诊断路径", ["twin_profiles", "twin_profile_nodes", "diagnosis_reports", "learning_plans"], PALETTE["purple"], PALETTE["purple_border"]),
        ("教师与互动", ["intervention_packages", "teaching_interaction_events", "homework_grading_events", "teaching_research_records"], PALETTE["green"], PALETTE["green_border"]),
    ]
    positions = [(80, 200), (1140, 200), (80, 690), (1140, 690)]
    for (title, tables, fill, border), (x, y) in zip(groups, positions):
        draw.rounded_rectangle((x, y, x + 880, y + 370), radius=24, fill="#F8FAFC", outline="#CBD5E1", width=3)
        draw.text((x + 35, y + 30), title, font=FONT_SUBTITLE, fill=PALETTE["ink"])
        for idx, table in enumerate(tables):
            tx = x + 50 + (idx % 2) * 405
            ty = y + 110 + (idx // 2) * 115
            count = schema["tables"].get(table, {}).get("row_count", 0)
            _box(draw, (tx, ty, tx + 350, ty + 80), f"{table}\n{count} 行", fill, border, FONT_SMALL, max_chars=22)
    _arrow(draw, (960, 385), (1140, 385), label="学习记录")
    _arrow(draw, (1580, 570), (1580, 690), label="汇聚画像")
    _arrow(draw, (960, 875), (1140, 875), label="风险反馈")
    _elbow_arrow(draw, [(520, 690), (520, 620), (1580, 620), (1580, 690)], label="证据消费")
    return _save(image, "fig_4_database_domains")


def fig_course_publish_flow() -> Path:
    image, draw = _canvas("图 5  课程底座建设、资源绑定与能力映射发布流程", "未发布草稿不得被学生端、诊断、路径或看板读取")
    steps = [
        ("教师录入课程结构", PALETTE["blue"], PALETTE["blue_border"]),
        ("生成课程知识图谱草稿", PALETTE["green"], PALETTE["green_border"]),
        ("绑定本地与外部资源候选", PALETTE["green"], PALETTE["green_border"]),
        ("行业情报生成岗位能力候选", PALETTE["cyan"], PALETTE["cyan_border"]),
        ("教师审核资源与能力映射", PALETTE["orange"], PALETTE["orange_border"]),
        ("发布课程底座版本", PALETTE["purple"], PALETTE["purple_border"]),
        ("下游只读已发布底座", PALETTE["purple"], PALETTE["purple_border"]),
    ]
    y = 170
    prev = None
    main_boxes = []
    for idx, (text, fill, border) in enumerate(steps):
        x1, x2 = (720, 1480)
        box = (x1, y, x2, y + 105)
        main_boxes.append(box)
        _box(draw, box, text, fill, border, max_chars=24)
        if prev:
            _arrow(draw, ((prev[0] + prev[2]) // 2, prev[3]), ((box[0] + box[2]) // 2, box[1]))
        prev = box
        y += 135
    gap_box = (90, 575, 570, 695)
    repair_box = (1620, 430, 2100, 550)
    _box(draw, gap_box, "能力缺口\n补充知识点建议", PALETTE["red"], PALETTE["red_border"], max_chars=16)
    _box(draw, repair_box, "资源不可用/映射不足\n退回教师修正", PALETTE["red"], PALETTE["red_border"], max_chars=18)
    resource_box = main_boxes[2]
    ability_box = main_boxes[3]
    graph_box = main_boxes[1]
    _arrow(draw, (ability_box[0], ability_box[1] + 52), (gap_box[2], gap_box[1] + 60))
    _arrow(draw, (resource_box[2], resource_box[1] + 52), (repair_box[0], repair_box[1] + 60))
    _elbow_arrow(draw, [(gap_box[2], gap_box[1] + 95), (635, gap_box[1] + 95), (635, graph_box[1] + 52), (graph_box[0], graph_box[1] + 52)])
    _elbow_arrow(draw, [(repair_box[0], repair_box[1] + 95), (1535, repair_box[1] + 95), (1535, resource_box[1] + 52), (resource_box[2], resource_box[1] + 52)])
    return _save(image, "fig_5_course_publish_flow")


def fig_student_diagnosis_path() -> Path:
    image, draw = _canvas("图 6  学生画像、诊断智能体与个性化路径协同流程", "诊断提供计算服务，学生端生成并执行补学路径")
    lanes = [
        ("证据输入", 160, PALETTE["orange"], PALETTE["orange_border"], ["测验作答", "作业/代码题", "资源学习", "5E 互动"]),
        ("学生数字孪生", 640, PALETTE["purple"], PALETTE["purple_border"], ["知识点掌握", "章节实践能力", "职业能力达成", "画像快照"]),
        ("诊断智能体", 1120, PALETTE["green"], PALETTE["green_border"], ["证据充分度", "薄弱点解释", "原因与置信度", "补证建议"]),
        ("个性化路径", 1600, PALETTE["blue"], PALETTE["blue_border"], ["路径版本", "资源/测验/作业入口", "节点状态", "完成回流"]),
    ]
    for title, x, fill, border, nodes in lanes:
        draw.rounded_rectangle((x, 180, x + 360, 1050), radius=24, fill="#F8FAFC", outline="#CBD5E1", width=3)
        draw.text((x + 36, 225), title, font=FONT_SUBTITLE, fill=PALETTE["ink"])
        for i, node in enumerate(nodes):
            y = 330 + i * 155
            _box(draw, (x + 40, y, x + 320, y + 86), node, fill, border, FONT_SMALL, max_chars=13)
    for x in (520, 1000, 1480):
        _arrow(draw, (x, 610), (x + 120, 610), label=["汇聚", "诊断", "生成"][int((x - 520) / 480)])
    _elbow_arrow(draw, [(1780, 950), (1780, 1130), (340, 1130), (340, 950)], label="学习结果回流")
    return _save(image, "fig_6_student_diagnosis_path")


def fig_teacher_intervention_loop() -> Path:
    image, draw = _canvas("图 7  教师看板、诊断证据与干预任务包闭环", "教师端只进行干预审核和任务下发，不替学生生成正式补学路径")
    nodes = [
        ((120, 250, 500, 370), "教师看板\n班级风险/学生名单", PALETTE["blue"], PALETTE["blue_border"]),
        ((650, 250, 1030, 370), "统一诊断服务\n证据与建议动作", PALETTE["green"], PALETTE["green_border"]),
        ((1180, 250, 1560, 370), "干预草稿\n资源/作业/提醒", PALETTE["orange"], PALETTE["orange_border"]),
        ((1710, 250, 2090, 370), "教师审核后下发", PALETTE["blue"], PALETTE["blue_border"]),
        ((1710, 600, 2090, 720), "学生执行任务包", PALETTE["purple"], PALETTE["purple_border"]),
        ((1180, 600, 1560, 720), "完成与评分记录", PALETTE["purple"], PALETTE["purple_border"]),
        ((650, 600, 1030, 720), "画像/看板/教师画像回流", PALETTE["cyan"], PALETTE["cyan_border"]),
        ((120, 600, 500, 720), "二次观察\n风险是否下降", PALETTE["blue"], PALETTE["blue_border"]),
    ]
    for xy, text, fill, border in nodes:
        _box(draw, xy, text, fill, border, max_chars=16)
    arrows = [((500, 310), (650, 310)), ((1030, 310), (1180, 310)), ((1560, 310), (1710, 310)), ((1900, 370), (1900, 600)), ((1710, 660), (1560, 660)), ((1180, 660), (1030, 660)), ((650, 660), (500, 660))]
    for a, b in arrows:
        _arrow(draw, a, b)
    _elbow_arrow(draw, [(310, 600), (310, 510), (310, 370)], label="继续监测")
    return _save(image, "fig_7_teacher_intervention_loop")


def fig_api_classification(operations: list[dict[str, Any]]) -> Path:
    grouped: dict[str, int] = defaultdict(int)
    for op in operations:
        grouped[classify_operation(op)] += 1
    items = sorted(grouped.items(), key=lambda item: item[1], reverse=True)[:14]
    image, draw = _canvas("图 8  FastAPI 接口分组统计", f"OpenAPI 当前实现共 {len(operations)} 个 operation")
    max_count = max(count for _, count in items)
    y = 190
    for module, count in items:
        bar_w = int(1180 * count / max_count)
        draw.text((120, y + 12), module, font=FONT_SMALL, fill=PALETTE["ink"])
        draw.rounded_rectangle((620, y, 620 + bar_w, y + 42), radius=12, fill="#DDEBFF", outline="#2F66B3", width=2)
        draw.text((640 + bar_w, y + 8), str(count), font=FONT_SMALL, fill=PALETTE["muted"])
        y += 68
    _box(draw, (1660, 260, 2080, 440), "说明\n接口归类采用业务主责\n而非简单 URL 前缀", PALETTE["green"], PALETTE["green_border"], max_chars=17)
    _box(draw, (1660, 560, 2080, 740), "支撑接口\nLLM/OCR/日志/资源文件\n单独列示", PALETTE["orange"], PALETTE["orange_border"], max_chars=17)
    return _save(image, "fig_8_api_classification")


def fig_gap_roadmap() -> Path:
    image, draw = _canvas("图 9  需求—实现差距与后续验收路线", "把“有接口/有表”继续推进为“边界闭环、页面可用、证据可追溯”")
    stages = [
        ("接口与表快照", "已完成\n164 接口 / 42 表", PALETTE["green"], PALETTE["green_border"]),
        ("边界代码验收", "发布态过滤\n教师授权\n确认标记", PALETTE["orange"], PALETTE["orange_border"]),
        ("页面真实流程", "zyh/teacher\n端到端验收", PALETTE["blue"], PALETTE["blue_border"]),
        ("图件与 Word", "详细设计\n正式交付版", PALETTE["purple"], PALETTE["purple_border"]),
        ("剩余功能补齐", "5E 有效性\n教研互动\n回流闭环", PALETTE["red"], PALETTE["red_border"]),
    ]
    x = 100
    prev = None
    for title, body, fill, border in stages:
        _box(draw, (x, 380, x + 330, 560), f"{title}\n{body}", fill, border, max_chars=14)
        if prev:
            _arrow(draw, (prev + 330, 470), (x, 470))
        prev = x
        x += 420
    notes = [
        "课程底座：只读 published",
        "作业覆盖：只采纳教师确认",
        "路径生成：学生端触发",
        "看板访问：教师授权过滤",
        "5E 证据：辅助而非替代",
    ]
    for i, note in enumerate(notes):
        _box(draw, (190 + i * 390, 760, 500 + i * 390, 850), note, PALETTE["gray"], PALETTE["gray_border"], FONT_SMALL, max_chars=14)
    return _save(image, "fig_9_gap_roadmap")


def generate_figures(operations: list[dict[str, Any]], schema: dict[str, Any]) -> list[tuple[str, Path]]:
    return [
        ("图 1 AI-Education 运行架构与边界", fig_system_architecture()),
        ("图 2 需求模块与实现模块协同关系", fig_module_collaboration()),
        ("图 3 课程底座—学习证据—画像诊断—教学反馈数据闭环", fig_data_closed_loop()),
        ("图 4 MySQL 数据库分域与表结构现状", fig_database_domains(schema)),
        ("图 5 课程底座建设、资源绑定与能力映射发布流程", fig_course_publish_flow()),
        ("图 6 学生画像、诊断智能体与个性化路径协同流程", fig_student_diagnosis_path()),
        ("图 7 教师看板、诊断证据与干预任务包闭环", fig_teacher_intervention_loop()),
        ("图 8 FastAPI 接口分组统计", fig_api_classification(operations)),
        ("图 9 需求—实现差距与后续验收路线", fig_gap_roadmap()),
    ]


def _drawio_node(
    node_id: str,
    value: str,
    x: int,
    y: int,
    width: int,
    height: int,
    fill: str,
    stroke: str,
) -> str:
    style = (
        "rounded=1;whiteSpace=wrap;html=1;arcSize=12;"
        f"fillColor={fill};strokeColor={stroke};strokeWidth=3;"
        "fontFamily=Microsoft YaHei;fontSize=26;fontColor=#172033;"
        "align=center;verticalAlign=middle;spacing=10;"
    )
    return (
        f'<mxCell id="{node_id}" value="{html.escape(value).replace(chr(10), "&lt;br&gt;")}" '
        f'style="{style}" vertex="1" parent="1">'
        f'<mxGeometry x="{x}" y="{y}" width="{width}" height="{height}" as="geometry"/>'
        "</mxCell>"
    )


def _drawio_edge(edge_id: str, source: str, target: str, label: str = "") -> str:
    style = (
        "edgeStyle=orthogonalEdgeStyle;rounded=0;orthogonalLoop=1;jettySize=auto;html=1;"
        "endArrow=block;endFill=1;strokeWidth=3;strokeColor=#42526E;"
        "fontFamily=Microsoft YaHei;fontSize=22;fontColor=#52606D;"
    )
    return (
        f'<mxCell id="{edge_id}" value="{html.escape(label)}" style="{style}" edge="1" parent="1" '
        f'source="{source}" target="{target}"><mxGeometry relative="1" as="geometry"/></mxCell>'
    )


def _write_drawio(name: str, title: str, nodes: list[dict[str, Any]], edges: list[tuple[str, str, str]]) -> Path:
    DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    cells = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>']
    cells.append(
        _drawio_node(
            "title",
            title,
            40,
            30,
            2080,
            70,
            "#F8FAFC",
            "#CBD5E1",
        )
    )
    for index, node in enumerate(nodes, start=1):
        cells.append(
            _drawio_node(
                f"n{index}",
                node["text"],
                node["x"],
                node["y"],
                node["w"],
                node["h"],
                node["fill"],
                node["stroke"],
            )
        )
    for index, (source, target, label) in enumerate(edges, start=1):
        cells.append(_drawio_edge(f"e{index}", source, target, label))
    xml = (
        '<mxfile host="app.diagrams.net" modified="2026-06-29T00:00:00.000Z" agent="Codex" version="24.7.17">'
        f'<diagram id="{name}" name="{html.escape(title[:40])}">'
        '<mxGraphModel dx="2200" dy="1300" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" '
        'arrows="1" fold="1" page="1" pageScale="1" pageWidth="2200" pageHeight="1300" math="0" shadow="0">'
        "<root>"
        + "".join(cells)
        + "</root></mxGraphModel></diagram></mxfile>"
    )
    path = DRAWIO_DIR / f"{name}.drawio"
    path.write_text(xml, encoding="utf-8")
    return path


def generate_drawio_sources() -> list[Path]:
    """Create editable diagrams.net sources aligned with the generated PNG figure set."""
    specs: list[tuple[str, str, list[dict[str, Any]], list[tuple[str, str, str]]]] = []

    specs.append(
        (
            "fig_1_system_architecture",
            "图 1 AI-Education 运行架构与边界",
            [
                {"text": "学生端\n学习、画像、路径、互动", "x": 80, "y": 190, "w": 540, "h": 120, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "教师端\n课程建设、看板、干预", "x": 80, "y": 390, "w": 540, "h": 120, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "管理员端\n账号、课程、运行治理", "x": 80, "y": 590, "w": 540, "h": 120, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "FastAPI 接口层\n164 个 OpenAPI operation", "x": 780, "y": 190, "w": 700, "h": 120, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "业务服务\n课程/测验/作业/路径", "x": 780, "y": 430, "w": 300, "h": 140, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "智能服务\n诊断/5E/行业情报", "x": 1180, "y": 430, "w": 300, "h": 140, "fill": PALETTE["purple"], "stroke": PALETTE["purple_border"]},
                {"text": "MySQL ai_education_design\n42 张真实业务与支撑表", "x": 780, "y": 690, "w": 700, "h": 130, "fill": PALETTE["orange"], "stroke": PALETTE["orange_border"]},
                {"text": "外部资源\nB站/YouTube/CSDN/文档", "x": 1640, "y": 260, "w": 440, "h": 140, "fill": PALETTE["cyan"], "stroke": PALETTE["cyan_border"]},
                {"text": "大模型与工具\nLLM/OCR/RAG/日志", "x": 1640, "y": 560, "w": 440, "h": 140, "fill": PALETTE["cyan"], "stroke": PALETTE["cyan_border"]},
            ],
            [("n1", "n4", ""), ("n2", "n4", ""), ("n3", "n4", ""), ("n4", "n5", ""), ("n4", "n6", ""), ("n5", "n7", ""), ("n6", "n7", ""), ("n5", "n8", ""), ("n6", "n9", "")],
        )
    )

    specs.append(
        (
            "fig_2_module_collaboration",
            "图 2 需求模块与实现模块协同关系",
            [
                {"text": "3.1 课程数字孪生", "x": 120, "y": 320, "w": 340, "h": 90, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "3.12 行业情报", "x": 120, "y": 500, "w": 340, "h": 90, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "3.2 学习空间", "x": 650, "y": 320, "w": 340, "h": 90, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "3.3 在线测验\n3.8 作业评测\n3.7 5E 辅助", "x": 650, "y": 500, "w": 340, "h": 150, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "3.4 学生数字孪生", "x": 1180, "y": 320, "w": 340, "h": 90, "fill": PALETTE["purple"], "stroke": PALETTE["purple_border"]},
                {"text": "3.5 诊断智能体\n3.6 个性化路径", "x": 1180, "y": 500, "w": 340, "h": 130, "fill": PALETTE["purple"], "stroke": PALETTE["purple_border"]},
                {"text": "3.10 教师看板\n3.9 教师干预\n3.11 教学互动", "x": 1710, "y": 390, "w": 340, "h": 160, "fill": PALETTE["orange"], "stroke": PALETTE["orange_border"]},
            ],
            [("n1", "n3", "已发布底座"), ("n2", "n1", "能力候选"), ("n3", "n5", "学习证据"), ("n4", "n5", "证据回流"), ("n5", "n6", "画像状态"), ("n6", "n7", "风险与建议"), ("n7", "n1", "反馈优化")],
        )
    )

    specs.append(
        (
            "fig_3_data_closed_loop",
            "图 3 课程底座—学习证据—画像诊断—教学反馈数据闭环",
            [
                {"text": "教师建设课程底座", "x": 120, "y": 230, "w": 320, "h": 110, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "发布课程图谱\n资源与能力映射", "x": 560, "y": 230, "w": 320, "h": 110, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "学生学习\n资源/测验/作业/互动", "x": 1000, "y": 230, "w": 320, "h": 110, "fill": PALETTE["orange"], "stroke": PALETTE["orange_border"]},
                {"text": "学生数字孪生\n画像与快照", "x": 1440, "y": 230, "w": 320, "h": 110, "fill": PALETTE["purple"], "stroke": PALETTE["purple_border"]},
                {"text": "诊断智能体\n薄弱点与原因", "x": 1440, "y": 520, "w": 320, "h": 110, "fill": PALETTE["purple"], "stroke": PALETTE["purple_border"]},
                {"text": "学生个性化路径\n补学与执行", "x": 1000, "y": 520, "w": 320, "h": 110, "fill": PALETTE["green"], "stroke": PALETTE["green_border"]},
                {"text": "教师看板\n班级风险与证据", "x": 560, "y": 520, "w": 320, "h": 110, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "教师干预任务包\n审核后下发", "x": 120, "y": 520, "w": 320, "h": 110, "fill": PALETTE["blue"], "stroke": PALETTE["blue_border"]},
                {"text": "结果回流\n画像更新 / 路径调整 / 课程维护 / 教师画像事件", "x": 560, "y": 820, "w": 760, "h": 130, "fill": PALETTE["cyan"], "stroke": PALETTE["cyan_border"]},
            ],
            [("n1", "n2", ""), ("n2", "n3", ""), ("n3", "n4", ""), ("n4", "n5", ""), ("n5", "n6", ""), ("n5", "n7", ""), ("n7", "n8", ""), ("n8", "n9", ""), ("n9", "n4", "证据更新"), ("n9", "n1", "教学反馈")],
        )
    )

    simple_specs = [
        (
            "fig_4_database_domains",
            "图 4 MySQL 数据库分域与表结构现状",
            ["课程与资源", "学习证据", "画像诊断路径", "教师与互动"],
        ),
        (
            "fig_5_course_publish_flow",
            "图 5 课程底座建设、资源绑定与能力映射发布流程",
            ["教师录入课程结构", "生成课程图谱草稿", "绑定资源候选", "生成岗位能力候选", "教师审核映射", "发布课程底座", "下游只读已发布"],
        ),
        (
            "fig_6_student_diagnosis_path",
            "图 6 学生画像、诊断智能体与个性化路径协同流程",
            ["证据输入", "学生数字孪生", "诊断智能体", "个性化路径", "学习结果回流"],
        ),
        (
            "fig_7_teacher_intervention_loop",
            "图 7 教师看板、诊断证据与干预任务包闭环",
            ["教师看板", "统一诊断服务", "干预草稿", "教师审核下发", "学生执行", "完成评分", "画像看板回流"],
        ),
        (
            "fig_8_api_classification",
            "图 8 FastAPI 接口分组统计",
            ["业务模块接口", "支撑接口", "前端兼容接口", "OpenAPI 164 operations"],
        ),
        (
            "fig_9_gap_roadmap",
            "图 9 需求—实现差距与后续验收路线",
            ["接口与表快照", "边界代码验收", "页面真实流程", "图件与 Word", "剩余功能补齐"],
        ),
    ]
    for name, title, labels in simple_specs:
        nodes = []
        for index, label in enumerate(labels):
            nodes.append(
                {
                    "text": label,
                    "x": 120 + index * 380 if len(labels) <= 5 else 180 + (index % 4) * 460,
                    "y": 430 if len(labels) <= 5 else 300 + (index // 4) * 250,
                    "w": 300,
                    "h": 110,
                    "fill": [PALETTE["blue"], PALETTE["green"], PALETTE["orange"], PALETTE["purple"], PALETTE["cyan"]][index % 5],
                    "stroke": [PALETTE["blue_border"], PALETTE["green_border"], PALETTE["orange_border"], PALETTE["purple_border"], PALETTE["cyan_border"]][index % 5],
                }
            )
        edges = [(f"n{i}", f"n{i + 1}", "") for i in range(1, len(labels))]
        specs.append((name, title, nodes, edges))

    return [_write_drawio(name, title, nodes, edges) for name, title, nodes, edges in specs]


def _md_table(headers: list[str], rows: list[list[Any]]) -> list[str]:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(item).replace("|", "\\|").replace("\n", " ") for item in row) + " |")
    return lines


def build_markdown(operations: list[dict[str, Any]], schema: dict[str, Any], figures: list[tuple[str, Path]]) -> str:
    grouped_ops: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for op in operations:
        grouped_ops[classify_operation(op)].append(op)
    tables = schema["tables"]
    lines: list[str] = [
        "# AI-Education 系统详细设计说明书（实现对齐高规格版）",
        "",
        f"生成时间：{datetime.now():%Y-%m-%d %H:%M:%S}",
        "",
        "## 1. 文档概述",
        "",
        "### 1.1 编写目的",
        "本文档用于在需求分析文档基础上，进一步明确 AI-Education 系统的模块设计、接口边界、数据库结构、关键流程、图件说明和当前实现差距。本文档以当前 FastAPI OpenAPI、真实 MySQL 数据库和已确认需求边界为依据，避免仅按需求文字或旧基线文档臆写。",
        "",
        "### 1.2 编制依据",
        f"- 需求基准：{REQ_DOC}",
        f"- 开发基线：{BASELINE_MD}",
        f"- 接口依据：当前 FastAPI `app.openapi()`，共 {len(operations)} 个 operation。",
        f"- 数据库依据：当前 MySQL `{schema['database']}`，共 {len(tables)} 张表。",
        "- 数据行数说明：本文出现的行数只代表本地演示与验证快照，不代表正式生产数据。",
        "",
        "### 1.3 设计原则",
        "系统设计坚持“课程底座先发布、学习证据可追溯、诊断服务不越界、学生路径由学生端执行、教师干预需审核、教师看板需授权过滤”的原则。所有 AI 生成结果在进入正式教学动作前，都必须保留人工确认或人工审核边界。",
        "",
        "## 2. 总体架构设计",
        "",
    ]
    for caption, path in figures[:4]:
        lines.append(f"![{caption}]({path.as_posix()})")
        lines.append(caption)
        lines.append("")
    lines.extend(
        [
            "系统采用前后端分离架构，前端负责学生端、教师端和管理员端页面组织；后端以 FastAPI 暴露业务接口；智能服务提供诊断、5E 引导、资源推荐和行业情报能力；MySQL 保存课程底座、学习证据、画像、诊断、路径、作业、干预和教师行为事件。",
            "",
            "模块协同关系不按历史 URL 前缀机械划分，而按业务主责划分。例如能力映射相关接口虽然挂在 `/api/course-digital-twin` 下，但在详细设计中主归属为行业情报与能力对接；资源学习事件作为学习行为证据支撑接口单列；LLM、OCR、日志和文件管理接口作为支撑能力单列。",
            "",
            "## 3. 核心业务流程设计",
            "",
        ]
    )
    for caption, path in figures[4:7]:
        lines.append(f"![{caption}]({path.as_posix()})")
        lines.append(caption)
        lines.append("")
    lines.extend(
        [
            "课程发布流程强调草稿与已发布底座隔离。课程结构、候选资源、岗位能力候选和能力映射在发布前均属于草稿或待审状态，下游学生端、诊断、路径和看板只能读取已发布结果。",
            "",
            "学生画像与诊断路径流程强调职责边界。学生数字孪生维护画像状态；诊断智能体提供薄弱点、原因和证据等级；个性化路径由学生端生成和执行；教师端不替学生生成正式补学路径，而是通过干预任务包安排资源、作业、提醒和跟踪。",
            "",
            "教师干预流程强调人工审核。诊断服务给出风险和建议动作，系统可生成干预草稿，但任务包必须经教师审核后下发。学生完成后，任务结果回流学生画像、教师看板和教师数字孪生。",
            "",
            "## 4. 模块详细设计",
            "",
        ]
    )
    module_rows = []
    for module, config in MODULES.items():
        ops = sorted(grouped_ops.get(module, []), key=lambda item: (item["path"], item["method"]))
        module_tables = [table for table in config["tables"] if table in tables]
        empty = [table for table in module_tables if tables[table]["row_count"] == 0]
        status = "基础实现，部分表为空：" + "、".join(empty) if empty else "已有接口和本地数据，需继续按页面和权限验收"
        module_rows.append([config["requirement"], module, config["note"], len(ops), "、".join(module_tables), status])
    lines.extend(_md_table(["章节", "模块", "职责边界", "接口数", "相关表", "当前状态"], module_rows))
    lines.append("")
    for module, config in MODULES.items():
        ops = sorted(grouped_ops.get(module, []), key=lambda item: (item["path"], item["method"]))
        module_tables = [table for table in config["tables"] if table in tables]
        lines.extend([f"### {config['requirement']} {module}", "", config["note"], ""])
        lines.extend(_md_table(["设计项", "当前说明"], [
            ["接口覆盖", f"{len(ops)} 个 operation"],
            ["数据表", "、".join(f"{table}({tables[table]['row_count']}行)" for table in module_tables)],
            ["关键边界", _module_boundary(module)],
            ["待验收点", "；".join(MODULE_GAPS[module])],
        ]))
        lines.append("")
    lines.extend(
        [
            "## 5. 接口设计",
            "",
            "接口章节以当前 FastAPI OpenAPI 为准。正式接口说明后续还应继续补请求字段、响应字段、权限角色和错误码；本版重点解决接口覆盖、模块归属和支撑接口分组。",
            "",
        ]
    )
    lines.append(f"![{figures[7][0]}]({figures[7][1].as_posix()})")
    lines.append(figures[7][0])
    lines.append("")
    for module in list(MODULES) + sorted(k for k in grouped_ops if k not in MODULES):
        ops = sorted(grouped_ops.get(module, []), key=lambda item: (item["path"], item["method"]))
        if not ops:
            continue
        lines.extend([f"### 5.{len([line for line in lines if line.startswith('### 5.')]) + 1} {module}", ""])
        rows = [[op["method"], op["path"], op["summary"], op["request_body"], op["responses"]] for op in ops]
        lines.extend(_md_table(["方法", "路径", "摘要", "请求体", "响应码"], rows[:30]))
        if len(rows) > 30:
            lines.append(f"本组接口共 {len(rows)} 个，表中展示前 30 个，完整清单见 generated/fastapi接口清单-当前实现.md。")
        lines.append("")
    lines.extend(
        [
            "## 6. 数据库设计",
            "",
            f"当前数据库为 `{schema['database']}`，共 {len(tables)} 张表。表结构来自 `SHOW FULL COLUMNS` 与 `SHOW INDEX`，正式数据库设计还应继续补充 `SHOW CREATE TABLE` 或 `information_schema.KEY_COLUMN_USAGE`，以区分普通索引、唯一约束和真实外键。",
            "",
        ]
    )
    table_rows = []
    for table, info in tables.items():
        table_rows.append([table, info["row_count"], len(info["columns"]), "、".join(classify_table(table)), "结构已建未产品化" if info["row_count"] == 0 else "已有本地数据"])
    lines.extend(_md_table(["表名", "当前行数", "字段数", "归属/消费模块", "数据状态"], table_rows))
    lines.append("")
    lines.extend(
        [
            "## 7. 需求实现差距与验收路线",
            "",
            f"![{figures[8][0]}]({figures[8][1].as_posix()})",
            figures[8][0],
            "",
            "当前系统已经具备主链路接口和数据库结构，但“有接口、有表、有本地数据”不等于业务边界已经闭环。后续验收必须重点检查发布态过滤、教师确认、学生/教师视图隔离、教师授权范围和证据回流。",
            "",
        ]
    )
    gap_rows = []
    for module, gaps in MODULE_GAPS.items():
        gap_rows.append([module, "；".join(gaps)])
    lines.extend(_md_table(["模块", "主要差距与验收点"], gap_rows))
    lines.append("")
    return "\n".join(lines) + "\n"


def _module_boundary(module: str) -> str:
    mapping = {
        "课程数字孪生与课程资源": "只发布课程底座、资源绑定和能力支撑关系；下游只读已发布版本。",
        "学生学习空间": "组织学生学习入口和学习行为记录，不重新计算画像或诊断。",
        "在线测验": "提供已发布知识点测验和作答记录，是知识点强证据来源。",
        "学生数字孪生": "维护学生画像状态，不解释根因，不生成路径。",
        "诊断智能体": "提供薄弱点、原因、证据等级和建议动作，不直接下发任务。",
        "个性化学习路径": "学生端生成和执行路径，正式节点来自已发布课程图谱。",
        "5E 教学智能体": "提供阶段化学习引导和过程性辅助证据。",
        "作业与实践评测": "作业覆盖知识点经教师确认后才影响叶子知识点画像。",
        "教师智能干预任务包": "教师审核后下发干预任务，学生完成后回流。",
        "教师看板与教师数字孪生": "查看授权范围内学情和教师行为画像，AI 建议需手动触发。",
        "教学互动": "人工沟通和教研记录层，行为事件回流教师数字孪生。",
        "行业情报与能力对接": "生成岗位和能力候选，不负责发布课程底座。",
    }
    return mapping.get(module, "")


def write_docx(markdown: str, figures: list[tuple[str, Path]]) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    fig_lookup = {path.as_posix(): (caption, path) for caption, path in figures}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)

    _setup_styles(doc)
    _setup_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(110)
    run = title.add_run("AI-Education 系统详细设计说明书")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(24)
    run.bold = True
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("实现对齐高规格版")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    r.bold = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"生成日期：{datetime.now():%Y-%m-%d}\n需求基准：{REQ_DOC.name}\n数据依据：FastAPI OpenAPI + MySQL ai_education_design")
    doc.add_page_break()

    doc.add_heading("修订说明", level=1)
    _add_table(doc, [["版本", "日期", "说明"], ["V1.0", f"{datetime.now():%Y-%m-%d}", "基于真实 FastAPI 接口和 MySQL 数据库生成高规格详细设计候选版"]], widths=[3, 4, 9])
    doc.add_paragraph("本文档为正式交付候选版，已补充统一风格图件。后续如系统接口或数据库发生变化，可重新运行生成脚本更新。")
    doc.add_page_break()

    doc.add_heading("目录", level=1)
    _add_toc(doc)
    doc.add_page_break()

    pending_table: list[list[str]] | None = None
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if line.startswith("# "):
            continue
        if not line:
            if pending_table:
                _add_table(doc, pending_table)
                pending_table = None
            continue
        if line.startswith("!["):
            if pending_table:
                _add_table(doc, pending_table)
                pending_table = None
            start = line.find("](")
            end = line.rfind(")")
            path_text = line[start + 2 : end]
            if path_text in fig_lookup:
                caption, path = fig_lookup[path_text]
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(path), width=Cm(16.2))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.style = doc.styles["Caption"]
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip(" `") for cell in line.strip("|").split("|")]
            if all(set(cell) <= {"-", ":"} for cell in cells):
                continue
            pending_table = pending_table or []
            pending_table.append(cells)
            continue
        if pending_table:
            _add_table(doc, pending_table)
            pending_table = None
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            _set_run_font(p)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(4)
            _set_run_font(p)
    if pending_table:
        _add_table(doc, pending_table)

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    doc.save(DOCX_PATH)


def _setup_styles(doc: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(23, 32, 51)
    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    caption.font.name = "宋体"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(82, 96, 109)


def _setup_header_footer(doc: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "AI-Education 系统详细设计说明书（实现对齐高规格版）"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(header, size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    _add_page_number(footer)
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)


def _add_page_number(paragraph: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def _add_toc(doc: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    p2 = doc.add_paragraph("打开 Word 后右键目录并选择“更新域”即可刷新页码。")
    _set_run_font(p2, size=9)


def _set_run_font(paragraph: Any, size: int | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if size:
            run.font.size = Pt(size)


def _add_table(doc: Any, rows: list[list[str]], widths: list[float] | None = None) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    if widths is None:
        widths = [16 / col_count] * col_count
    for row_idx, source_row in enumerate(rows):
        cells = table.rows[0].cells if row_idx == 0 else table.add_row().cells
        for idx in range(col_count):
            text = source_row[idx] if idx < len(source_row) else ""
            cell = cells[idx]
            cell.text = _shorten(text, 420 if col_count <= 4 else 220)
            if idx < len(widths):
                cell.width = Cm(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "D9EAF7" if row_idx == 0 else "FFFFFF")
            tc_pr.append(shd)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(8.5)
                    if row_idx == 0:
                        run.bold = True
    doc.add_paragraph()


def _shorten(text: str, limit: int) -> str:
    return text if len(text) <= limit else text[: limit - 1] + "…"


def main() -> None:
    HIGH_SPEC_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    DRAWIO_DIR.mkdir(parents=True, exist_ok=True)
    operations = load_openapi_operations()
    schema = load_database_schema()
    figures = generate_figures(operations, schema)
    drawio_paths = generate_drawio_sources()
    markdown = build_markdown(operations, schema, figures)
    MD_PATH.write_text(markdown, encoding="utf-8")
    write_docx(markdown, figures)
    summary = {
        "docx": str(DOCX_PATH),
        "markdown": str(MD_PATH),
        "figure_dir": str(FIG_DIR),
        "drawio_dir": str(DRAWIO_DIR),
        "figure_count": len(figures),
        "drawio_count": len(drawio_paths),
        "operation_count": len(operations),
        "table_count": len(schema["tables"]),
    }
    (HIGH_SPEC_DIR / "high_spec_detailed_design_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
