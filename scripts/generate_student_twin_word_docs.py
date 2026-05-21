from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "学生数字孪生Word文档"
IMG_DIR = OUT_DIR / "images"


TITLE_COLOR = RGBColor(31, 78, 121)
ACCENT = RGBColor(47, 117, 181)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, text, fill="#F7FBFF", outline="#2F75B5", text_color="#1F2937"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    fnt = font(28, True)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 32)
    line_h = 36
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=fnt, fill=text_color)
        y += line_h


def arrow(draw, start, end, color="#6B7280", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        points = [(x2, y2), (x2 - sign * 16, y2 - 9), (x2 - sign * 16, y2 + 9)]
    else:
        sign = 1 if y2 >= y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - sign * 16), (x2 + 9, y2 - sign * 16)]
    draw.polygon(points, fill=color)


def base_canvas(title: str, w=1500, h=900):
    img = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((50, 32), title, font=font(38, True), fill="#17365D")
    draw.line((50, 86, w - 50, 86), fill="#9DB7D5", width=3)
    return img, draw


def save_architecture(path: Path):
    img, d = base_canvas("学生数字孪生模块体系结构图")
    boxes = {
        "front": (90, 160, 360, 270, "学生诊断页面\nStudentTwinView"),
        "api": (610, 150, 890, 280, "数字孪生 API\n/api/digital-twin"),
        "collector": (180, 430, 480, 555, "DataCollector\n数据采集"),
        "calc": (600, 430, 900, 555, "ScoreCalculator\n掌握度计算"),
        "service": (1020, 430, 1320, 555, "StudentTwinService\n画像摘要"),
        "db": (540, 680, 960, 800, "实体化存储\nprofiles / nodes / history"),
        "path": (1080, 160, 1370, 280, "PathPlannerAgent\n学习路径推荐"),
    }
    for xy_text in boxes.values():
        draw_box(d, xy_text[:4], xy_text[4])
    arrow(d, (360, 215), (610, 215))
    arrow(d, (750, 280), (750, 430))
    arrow(d, (480, 492), (600, 492))
    arrow(d, (900, 492), (1020, 492))
    arrow(d, (750, 555), (750, 680))
    arrow(d, (960, 735), (1170, 555))
    arrow(d, (890, 215), (1080, 215))
    arrow(d, (1170, 280), (1170, 430))
    img.save(path)


def save_data_flow(path: Path):
    img, d = base_canvas("学生数字孪生数据流图")
    sources = [
        (80, 170, 330, 270, "学习进度\nprogress"),
        (80, 330, 330, 430, "测验成绩\nquiz_score"),
        (80, 490, 330, 590, "AI 交互\nllm_count"),
        (80, 650, 330, 750, "会话时长\nduration"),
    ]
    for item in sources:
        draw_box(d, item[:4], item[4], fill="#F2F7F2", outline="#548235")
        arrow(d, (330, (item[1] + item[3]) // 2), (520, 460))
    draw_box(d, (520, 360, 820, 560), "采集与归一化\nDataCollector")
    draw_box(d, (930, 360, 1230, 560), "掌握度计算\n0.4测验+0.3进度\n+0.2互动+0.1时长")
    draw_box(d, (930, 660, 1230, 790), "画像输出\n雷达/风险/趋势/薄弱点", fill="#FFF8EE", outline="#C55A11")
    arrow(d, (820, 460), (930, 460))
    arrow(d, (1080, 560), (1080, 660))
    img.save(path)


def save_er(path: Path):
    img, d = base_canvas("学生数字孪生数据库关系图")
    entities = [
        (70, 160, 330, 300, "users\nuser_id\nusername\nuser_type"),
        (520, 160, 830, 300, "twin_profiles\nusername\nuser_id\noverall_mastery"),
        (1010, 160, 1360, 300, "twin_profile_nodes\nusername + node_id\nmastery_score"),
        (520, 420, 830, 560, "twin_history\nusername + date\noverall_mastery"),
        (70, 420, 330, 560, "sessions\nsession_id\ncurrent_node"),
        (1010, 420, 1360, 560, "quiz_attempts\nattempt_id\nscore / total"),
        (520, 675, 830, 815, "course_nodes\ncourse_id + node_id\nnode_path"),
    ]
    for item in entities:
        draw_box(d, item[:4], item[4], fill="#F8FBFF", outline="#2F75B5")
    arrow(d, (330, 230), (520, 230))
    arrow(d, (830, 230), (1010, 230))
    arrow(d, (675, 300), (675, 420))
    arrow(d, (330, 490), (520, 490))
    arrow(d, (1185, 420), (1185, 300))
    arrow(d, (675, 675), (1010, 300))
    img.save(path)


def save_ui(path: Path):
    img, d = base_canvas("学生诊断页面界面结构草图", w=1500, h=950)
    draw_box(d, (80, 130, 1420, 240), "页头：学习诊断报告 / 重新生成诊断按钮", fill="#EAF3F8", outline="#2F75B5")
    draw_box(d, (80, 275, 1420, 345), "元信息栏：生成时间 / 风险等级 / 薄弱知识点 / 优势知识点", fill="#F7F7F7", outline="#A6A6A6")
    draw_box(d, (80, 390, 710, 620), "能力画像雷达图\n知识掌握、学习投入、实践能力、稳定性、测验表现", fill="#F3F8FF", outline="#4472C4")
    draw_box(d, (790, 390, 1420, 620), "学习趋势折线图\n近 30 天整体掌握度变化", fill="#F3F8FF", outline="#4472C4")
    draw_box(d, (80, 675, 710, 875), "学习风险预警列表\n高/中/低风险与建议", fill="#FFF2F2", outline="#C00000")
    draw_box(d, (790, 675, 1420, 875), "薄弱知识点列表\n掌握度、进度、测验分、节点路径", fill="#FFF8EE", outline="#C55A11")
    img.save(path)


def save_core_flow(path: Path):
    img, d = base_canvas("核心算法与服务调用流程")
    steps = [
        (80, 180, 330, 290, "接收学生标识"),
        (430, 180, 680, 290, "加载画像快照"),
        (780, 180, 1030, 290, "读取30天趋势"),
        (1130, 180, 1380, 290, "生成摘要JSON"),
        (210, 480, 510, 610, "计算 mastery_score"),
        (600, 480, 900, 610, "判定技术层级"),
        (990, 480, 1290, 610, "识别风险与薄弱点"),
    ]
    for s in steps:
        draw_box(d, s[:4], s[4])
    for a, b in [((330, 235), (430, 235)), ((680, 235), (780, 235)), ((1030, 235), (1130, 235))]:
        arrow(d, a, b)
    arrow(d, (555, 290), (360, 480))
    arrow(d, (555, 290), (750, 480))
    arrow(d, (905, 290), (1140, 480))
    d.text((180, 720), "mastery = 0.40×测验 + 0.30×进度 + 0.20×互动归一化 + 0.10×时长归一化", font=font(30, True), fill="#17365D")
    img.save(path)


def save_test(path: Path):
    img, d = base_canvas("学生数字孪生测试结构图")
    draw_box(d, (570, 150, 930, 260), "验收测试\n学生诊断闭环")
    draw_box(d, (450, 340, 1050, 450), "集成测试\nAPI + Store + 前端页面")
    draw_box(d, (300, 530, 1200, 640), "模块测试\n采集 / 计算 / 趋势 / 风险 / 路径")
    draw_box(d, (150, 720, 1350, 830), "单元测试\nScoreCalculator、StudentTwinService、TrendTracker、ProfileStore")
    arrow(d, (750, 720), (750, 640))
    arrow(d, (750, 530), (750, 450))
    arrow(d, (750, 340), (750, 260))
    img.save(path)


def save_plan(path: Path):
    img, d = base_canvas("学生数字孪生项目计划图")
    phases = [
        ("需求确认", 120, 180, 360),
        ("数据模型", 300, 280, 560),
        ("采集计算", 520, 380, 820),
        ("画像服务", 760, 480, 1040),
        ("前端联调", 980, 580, 1230),
        ("测试验收", 1160, 680, 1380),
    ]
    for i, (name, x1, y, x2) in enumerate(phases):
        d.rounded_rectangle((x1, y, x2, y + 72), radius=14, fill="#EAF3F8", outline="#2F75B5", width=3)
        d.text((x1 + 18, y + 18), name, font=font(26, True), fill="#17365D")
        if i:
            arrow(d, (phases[i - 1][3], phases[i - 1][2] + 36), (x1, y + 36))
    img.save(path)


def set_doc_style(doc: Document):
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = TITLE_COLOR
        style.font.bold = True


def shade_cell(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_cover(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Education 学生数字孪生模块")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = TITLE_COLOR
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = TITLE_COLOR
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(12)
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("版本", "V1.0"),
        ("日期", "2026-05-13"),
        ("模块", "学生数字孪生"),
        ("依据", "当前项目代码、接口、数据库与模板文档结构"),
    ]
    for row, (k, v) in zip(table.rows, data):
        row.cells[0].text = k
        row.cells[1].text = v
    doc.add_page_break()


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]):
    headers = list(headers)
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_picture(doc: Document, image: Path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Cm(width_cm))


def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, title, subtitle)
    return doc


def save(doc: Document, name: str):
    path = OUT_DIR / name
    doc.save(path)
    return path


def make_project_plan(images):
    doc = new_doc("附录C-2 项目计划", "面向学生数字孪生模块")
    doc.add_heading("1 项目概述", 1)
    doc.add_paragraph("学生数字孪生模块负责汇聚学习过程数据，生成学生画像、能力雷达、薄弱知识点、风险预警和趋势分析，并向课程孪生、教师孪生和个性化路径模块输出可复用结果。")
    add_picture(doc, images["plan"])
    doc.add_heading("2 项目范围", 1)
    add_table(doc, ["范围项", "说明"], [
        ("包含", "数据采集、画像计算、趋势记录、画像接口、学生诊断页面、路径推荐联动。"),
        ("不包含", "教师数字孪生完整指标体系、行业情报、作业中心全流程、生产级统一认证。"),
        ("依赖", "课程节点、测验记录、会话记录、LLM 日志、用户与师生关系数据。"),
    ])
    doc.add_heading("3 工作分解", 1)
    add_table(doc, ["阶段", "任务", "交付物"], [
        ("需求分析", "确定画像指标、输入数据和输出对象", "需求规格说明书"),
        ("数据设计", "设计画像主表、节点明细、趋势历史", "数据库设计报告"),
        ("核心实现", "实现采集、评分、风险、趋势和摘要生成", "核心技术实现报告"),
        ("界面联调", "学生诊断页面接入雷达图、趋势图和列表", "用户界面设计"),
        ("测试验收", "完成接口、算法、页面和回归测试", "系统测试报告"),
    ])
    doc.add_heading("4 风险计划", 1)
    add_table(doc, ["风险", "影响", "应对"], [
        ("历史 JSON 与实体表并存", "数据来源混淆", "以实体表为准，保留迁移兜底。"),
        ("画像数据不足", "风险与趋势不稳定", "缺失数据按 0 或稳定状态处理。"),
        ("LLM 日志无法匹配节点", "互动指标偏低", "用 topic、module 与叶子节点做模糊匹配。"),
        ("多课程逻辑未完全展开", "课程画像过滤不足", "保留 course_id 参数，逐步完善多课程映射。"),
    ])
    doc.add_heading("5 交付清单", 1)
    add_bullets(doc, [
        "学生画像数据模型与存储能力。",
        "掌握度计算、技术分层、风险预警和趋势分析逻辑。",
        "学生画像 API 与课程画像 API。",
        "学生端学习诊断报告页面。",
        "模块测试与验收说明。",
    ])
    return save(doc, "附录C-2 学生数字孪生项目计划.docx")


def make_requirements(images):
    doc = new_doc("附录G-2 产品需求规格说明书", "学生数字孪生模块")
    doc.add_heading("1 产品定位", 1)
    doc.add_paragraph("学生数字孪生是 AI-Education 的学习诊断核心模块，目标是把分散的学习行为转换为可解释、可展示、可联动的学生能力画像。")
    add_picture(doc, images["data_flow"])
    doc.add_heading("2 用户与使用场景", 1)
    add_table(doc, ["用户", "场景", "核心诉求"], [
        ("学生", "查看学习诊断报告", "知道自己强在哪里、弱在哪里、下一步学什么。"),
        ("教师", "查看学生画像摘要", "识别风险学生和薄弱知识点，进行针对性干预。"),
        ("课程孪生", "消费学生课程画像", "获得课程维度掌握度、风险等级和节点弱项。"),
        ("路径推荐", "读取薄弱节点", "生成个性化学习顺序和资源推荐。"),
    ])
    doc.add_heading("3 功能需求", 1)
    add_table(doc, ["编号", "需求", "验收标准"], [
        ("ST-FR-01", "采集学习进度、测验分、AI 交互和会话时长。", "调用采集接口后画像节点数据更新。"),
        ("ST-FR-02", "按权重计算知识点掌握度和整体掌握度。", "输出范围为 0-100，保留两位小数。"),
        ("ST-FR-03", "生成五维能力雷达图数据。", "返回知识掌握、学习投入、实践能力、稳定性、测验表现。"),
        ("ST-FR-04", "识别薄弱知识点。", "掌握度低于 60 的节点按分数升序输出。"),
        ("ST-FR-05", "生成学习风险预警。", "输出风险编码、等级、标题和说明。"),
        ("ST-FR-06", "提供 30 天趋势分析。", "输出 upward、stable 或 downward。"),
        ("ST-FR-07", "提供学生课程画像查询。", "POST /api/digital-twin/student-course-profile 返回课程节点画像。"),
    ])
    doc.add_heading("4 非功能需求", 1)
    add_table(doc, ["类别", "要求"], [
        ("准确性", "计算公式固定，缺失测验分按 0 处理，结果限制在 0-100。"),
        ("可解释性", "风险、等级和弱项均返回文字说明。"),
        ("可扩展性", "画像字段使用 Pydantic 模型和实体表，支持后续扩维。"),
        ("性能", "画像读取走数据库实体表，趋势按最近 30 天过滤。"),
        ("兼容性", "兼容 username、login_id、user_id 等历史标识。"),
    ])
    doc.add_heading("5 数据字典摘要", 1)
    add_table(doc, ["字段", "含义", "来源"], [
        ("quiz_score", "知识点测验得分", "QuizModule / quiz_attempts"),
        ("progress", "知识点学习进度", "学习模块 / 用户进度"),
        ("llm_interaction_count", "围绕知识点的 AI 交互次数", "llm_logs"),
        ("study_duration_minutes", "知识点学习时长", "sessions"),
        ("mastery_score", "知识点综合掌握度", "ScoreCalculator"),
        ("overall_mastery", "整体掌握度", "所有节点掌握度平均值"),
    ])
    return save(doc, "附录G-2 学生数字孪生产品需求规格说明书.docx")


def make_architecture(images):
    doc = new_doc("附录I-1 体系结构设计报告", "学生数字孪生模块")
    doc.add_heading("1 架构视图", 1)
    doc.add_paragraph("模块采用 API 层、服务层、计算层、存储层分层设计。前端只消费画像摘要，后端负责采集、归一化、计算和持久化。")
    add_picture(doc, images["architecture"])
    doc.add_heading("2 模块职责", 1)
    add_table(doc, ["模块", "职责"], [
        ("digital_twin_api.py", "提供采集、画像查询、测验分更新、路径生成、课程画像等 HTTP 接口。"),
        ("DataCollector", "采集学习进度、LLM 交互、会话时长和测验成绩。"),
        ("ScoreCalculator", "计算知识点掌握度和整体掌握度。"),
        ("StudentTwinService", "生成雷达、风险、趋势、弱项和跨模块输出。"),
        ("TwinProfileStore", "读写画像主表和节点明细。"),
        ("TrendTracker", "记录和查询每日掌握度快照。"),
        ("PathPlannerAgent", "基于薄弱点生成学习路径和资源推荐。"),
    ])
    doc.add_heading("3 接口关系", 1)
    add_table(doc, ["接口", "方法", "说明"], [
        ("/api/digital-twin/collect/{username}", "POST", "触发采集并刷新画像。"),
        ("/api/digital-twin/profile/{username}", "GET", "返回原始画像快照。"),
        ("/api/digital-twin/student-profile/{username}", "GET", "返回学生诊断摘要。"),
        ("/api/digital-twin/quiz-score", "POST", "回写知识点测验分。"),
        ("/api/digital-twin/student-course-profile", "POST", "返回学生在指定课程下的画像。"),
    ])
    doc.add_heading("4 设计原则", 1)
    add_bullets(doc, [
        "实体表优先，历史文件仅作兼容来源。",
        "采集与计算分离，便于独立测试。",
        "输出 JSON 保持稳定，前端图表只依赖摘要字段。",
        "风险和等级规则显式编码，保证可解释。",
    ])
    return save(doc, "附录I-1 学生数字孪生体系结构设计报告.docx")


def make_ui(images):
    doc = new_doc("附录I-2 用户界面设计", "学生数字孪生学习诊断页面")
    doc.add_heading("1 页面目标", 1)
    doc.add_paragraph("学生诊断页面用于把复杂画像结果转化为可扫读的学习报告，重点展示生成时间、风险等级、能力雷达、学习趋势、风险预警和薄弱知识点。")
    add_picture(doc, images["ui"])
    doc.add_heading("2 信息架构", 1)
    add_table(doc, ["区域", "展示内容", "数据字段"], [
        ("报告头部", "标题、说明、重新生成诊断按钮", "loading、refreshStudentTwin"),
        ("元信息栏", "生成时间、风险等级、薄弱点数、优势点数", "generated_at、risk_alerts、node_summary"),
        ("能力画像", "五维雷达图", "radar[]"),
        ("学习趋势", "近 30 天折线图", "trend.points[]"),
        ("风险预警", "风险等级、标题、建议", "risk_alerts[]"),
        ("薄弱知识点", "节点、掌握度、进度、测验分、路径", "weak_nodes[]"),
    ])
    doc.add_heading("3 交互设计", 1)
    add_bullets(doc, [
        "进入页面后自动读取当前用户并加载画像摘要。",
        "画像不存在时尝试触发刷新后再次读取。",
        "点击“重新生成诊断”后调用采集接口并重绘图表。",
        "薄弱知识点分页显示，每页 5 条。",
        "窗口尺寸变化时雷达图和趋势图自动 resize。",
    ])
    doc.add_heading("4 可视化规则", 1)
    add_table(doc, ["组件", "规则"], [
        ("风险等级", "存在 high 风险显示高风险，存在 medium 风险显示中等风险，否则低风险。"),
        ("雷达图", "各维度最大值 100，面积代表综合能力结构。"),
        ("趋势图", "Y 轴范围 0-100，平滑折线展示掌握度变化。"),
        ("进度条", "薄弱节点掌握度按百分比填充。"),
    ])
    return save(doc, "附录I-2 学生数字孪生用户界面设计.docx")


def make_database(images):
    doc = new_doc("附录I-3 数据库设计报告", "学生数字孪生模块")
    doc.add_heading("1 数据库关系", 1)
    add_picture(doc, images["er"])
    doc.add_heading("2 核心表设计", 1)
    add_table(doc, ["表名", "主键", "用途"], [
        ("users", "user_id", "学生身份、登录标识和资料快照。"),
        ("twin_profiles", "username", "学生画像主表，保存整体掌握度和更新时间。"),
        ("twin_profile_nodes", "username + node_id", "知识点画像明细，保存进度、测验、互动、时长和掌握度。"),
        ("twin_history", "username + snapshot_date", "每日整体掌握度快照，用于趋势分析。"),
        ("sessions", "session_id", "会话、当前节点和 PDF，用于学习时长采集。"),
        ("llm_logs", "id", "AI 交互日志，用于互动次数采集。"),
        ("quiz_attempts", "attempt_id", "测验记录，用于测验成绩同步。"),
        ("course_nodes", "course_id + node_id", "课程节点和路径，用于节点匹配。"),
    ])
    doc.add_heading("3 twin_profile_nodes 字段", 1)
    add_table(doc, ["字段", "类型", "说明"], [
        ("username", "TEXT", "学生用户名。"),
        ("user_id", "INTEGER", "学生数字 ID。"),
        ("node_id", "TEXT", "课程节点 ID。"),
        ("node_path_json", "TEXT", "节点路径 JSON 数组。"),
        ("quiz_score", "REAL", "测验得分。"),
        ("progress", "REAL", "学习进度。"),
        ("study_duration_minutes", "REAL", "学习时长分钟数。"),
        ("llm_interaction_count", "INTEGER", "AI 交互次数。"),
        ("mastery_score", "REAL", "综合掌握度。"),
        ("updated_at", "TEXT", "更新时间。"),
    ])
    doc.add_heading("4 索引与一致性", 1)
    add_bullets(doc, [
        "twin_profiles.user_id 建索引，支持按用户 ID 查询画像。",
        "twin_profile_nodes.username 建索引，支持快速加载学生全部节点。",
        "twin_history.user_id 建索引，支持趋势历史查询。",
        "quiz_attempts.created_at 建索引，支持测验历史审计。",
        "课程、测验、画像同步以实体表为准，避免 JSON 双写造成冲突。",
    ])
    return save(doc, "附录I-3 学生数字孪生数据库设计报告.docx")


def make_core(images):
    doc = new_doc("附录I-4 核心技术实现报告", "学生数字孪生模块")
    doc.add_heading("1 技术路线", 1)
    doc.add_paragraph("模块采用“数据采集 -> 归一化计算 -> 画像摘要 -> 前端可视化 -> 路径联动”的技术路线。核心规则以代码显式实现，减少不可解释的黑盒判断。")
    add_picture(doc, images["core"])
    doc.add_heading("2 掌握度算法", 1)
    doc.add_paragraph("单个知识点掌握度计算公式如下：")
    doc.add_paragraph("mastery_score = 0.40 × quiz_score + 0.30 × progress + 0.20 × min(llm_interaction_count / 10, 1) × 100 + 0.10 × min(study_duration_minutes / 30, 1) × 100")
    add_table(doc, ["输入", "权重", "说明"], [
        ("quiz_score", "40%", "测验表现，缺失时按 0 处理。"),
        ("progress", "30%", "学习进度，反映课程推进程度。"),
        ("llm_interaction_count", "20%", "AI 互动次数，10 次封顶归一化。"),
        ("study_duration_minutes", "10%", "学习时长，30 分钟封顶归一化。"),
    ])
    doc.add_heading("3 能力分层", 1)
    add_table(doc, ["条件", "层级", "含义"], [
        ("overall_mastery < 40", "基础薄弱", "优先补齐基础知识。"),
        ("40 <= overall_mastery < 60", "基础建立中", "已有基础，但短板明显。"),
        ("60 <= overall_mastery < 80 且薄弱点 <= 4", "能力成型", "可进入系统强化训练。"),
        ("60 <= overall_mastery < 80 且薄弱点 > 4", "基础建立中", "整体中等但薄弱点较多。"),
        ("overall_mastery >= 80", "进阶提升", "可进入高阶任务和项目实践。"),
    ])
    doc.add_heading("4 风险规则", 1)
    add_table(doc, ["风险", "触发规则"], [
        ("知识薄弱风险", "overall_mastery < 45 为高风险，45-60 为中风险。"),
        ("进度滞后风险", "平均进度 < 50 为高风险，50-70 为中风险。"),
        ("学习投入风险", "学习投入分 < 45 触发中风险。"),
        ("趋势下滑风险", "30 天趋势状态为 downward。"),
        ("薄弱点集中风险", "薄弱节点数 >= 5。"),
    ])
    doc.add_heading("5 输出结构", 1)
    add_table(doc, ["输出", "用途"], [
        ("radar", "前端能力雷达图。"),
        ("weak_nodes", "薄弱知识点列表和路径推荐输入。"),
        ("risk_alerts", "学习风险预警。"),
        ("trend", "趋势折线图和教师观察依据。"),
        ("node_summary", "统计摘要。"),
        ("outputs.for_course_twin", "课程孪生消费字段。"),
        ("outputs.for_teacher_twin", "教师孪生消费字段。"),
    ])
    return save(doc, "附录I-4 学生数字孪生核心技术实现报告.docx")


def make_test(images):
    doc = new_doc("附录K-1 系统测试报告", "学生数字孪生模块")
    doc.add_heading("1 测试目标", 1)
    doc.add_paragraph("验证学生数字孪生从数据采集、画像计算、接口输出到前端展示的闭环正确性、稳定性和可解释性。")
    add_picture(doc, images["test"])
    doc.add_heading("2 测试范围", 1)
    add_table(doc, ["范围", "说明"], [
        ("单元测试", "掌握度计算、能力分层、风险识别、趋势判断。"),
        ("接口测试", "采集、画像查询、课程画像、测验分回写。"),
        ("集成测试", "测验完成后同步画像，前端刷新后图表更新。"),
        ("异常测试", "画像不存在、学生不存在、course_id 缺失、趋势数据不足。"),
        ("回归测试", "学生端诊断页面、学习进度、路径推荐相关流程。"),
    ])
    doc.add_heading("3 测试用例摘要", 1)
    add_table(doc, ["编号", "用例", "预期结果"], [
        ("ST-TC-01", "输入 quiz=100、progress=100、互动和时长为空。", "mastery_score 不低于 90。"),
        ("ST-TC-02", "整体掌握度 35。", "技术层级为基础薄弱，产生知识薄弱高风险。"),
        ("ST-TC-03", "近 30 天掌握度提升超过 5。", "趋势状态为 upward。"),
        ("ST-TC-04", "薄弱节点数达到 5。", "产生薄弱点集中风险。"),
        ("ST-TC-05", "调用 /student-profile/{username}。", "返回 radar、weak_nodes、risk_alerts、trend。"),
        ("ST-TC-06", "调用 /student-course-profile 且学生不存在。", "返回 404。"),
        ("ST-TC-07", "前端点击重新生成诊断。", "刷新接口执行，图表重新渲染。"),
    ])
    doc.add_heading("4 当前验证命令", 1)
    doc.add_paragraph("后端语法与接口自测：")
    doc.add_paragraph("python -m py_compile DigitalTwinModule\\digital_twin_api.py")
    doc.add_paragraph("python scripts\\backend_api_selftest.py --base-url http://127.0.0.1:8000")
    doc.add_paragraph("前端构建验证：")
    doc.add_paragraph("cd frontend-vue && npm run build")
    doc.add_heading("5 测试结论", 1)
    doc.add_paragraph("学生数字孪生模块测试重点明确，核心算法可用单元用例验证，接口可通过后端自测覆盖，前端页面可通过构建和人工交互验收。后续建议补充专门的 StudentTwinService 单元测试和 Playwright 页面截图测试。")
    return save(doc, "附录K-1 学生数字孪生系统测试报告.docx")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "architecture": IMG_DIR / "architecture.png",
        "data_flow": IMG_DIR / "data_flow.png",
        "er": IMG_DIR / "er.png",
        "ui": IMG_DIR / "ui.png",
        "core": IMG_DIR / "core_flow.png",
        "test": IMG_DIR / "test.png",
        "plan": IMG_DIR / "plan.png",
    }
    save_architecture(images["architecture"])
    save_data_flow(images["data_flow"])
    save_er(images["er"])
    save_ui(images["ui"])
    save_core_flow(images["core"])
    save_test(images["test"])
    save_plan(images["plan"])

    paths = [
        make_project_plan(images),
        make_requirements(images),
        make_architecture(images),
        make_ui(images),
        make_database(images),
        make_core(images),
        make_test(images),
    ]
    print("\n".join(str(path) for path in paths))


if __name__ == "__main__":
    main()

