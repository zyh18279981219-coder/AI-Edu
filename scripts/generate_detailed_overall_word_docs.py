from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "整体系统Word文档"
IMG_DIR = OUT_DIR / "images_detailed"
TITLE_COLOR = RGBColor(31, 78, 121)


def get_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, width: int) -> list[str]:
    lines: list[str] = []
    cur = ""
    for ch in text:
        t = cur + ch
        box = draw.textbbox((0, 0), t, font=fnt)
        if box[2] - box[0] <= width:
            cur = t
        else:
            if cur:
                lines.append(cur)
            cur = ch
    if cur:
        lines.append(cur)
    return lines


def image_base(title: str, w=1700, h=1000):
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((60, 34), title, font=get_font(42, True), fill="#17365D")
    d.line((60, 96, w - 60, 96), fill="#9DB7D5", width=3)
    return img, d


def draw_box(d, xy, text, fill="#F7FBFF", outline="#2F75B5", size=25):
    x1, y1, x2, y2 = xy
    d.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    fnt = get_font(size, True)
    lines = wrap(d, text, fnt, x2 - x1 - 30)
    line_h = int(size * 1.35)
    y = y1 + (y2 - y1 - line_h * len(lines)) / 2
    for line in lines:
        box = d.textbbox((0, 0), line, font=fnt)
        d.text((x1 + (x2 - x1 - (box[2] - box[0])) / 2, y), line, font=fnt, fill="#1F2937")
        y += line_h


def arrow(d, start, end, color="#6B7280", width=4):
    d.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - sign * 17, y2 - 10), (x2 - sign * 17, y2 + 10)]
    else:
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 10, y2 - sign * 17), (x2 + 10, y2 - sign * 17)]
    d.polygon(pts, fill=color)


def make_scope_image(path: Path):
    img, d = image_base("整体系统业务范围图")
    center = (710, 420, 990, 560)
    draw_box(d, center, "AI-Education\n智能教育平台", fill="#EAF3F8")
    items = [
        ("学生学习闭环", 110, 180, "#F2F7F2", "#548235"),
        ("教师教学闭环", 1240, 180, "#F2F7F2", "#548235"),
        ("管理员运维", 110, 650, "#F7F7F7", "#7F7F7F"),
        ("智能体与RAG", 1240, 650, "#FFF2F2", "#C00000"),
        ("数字孪生", 710, 170, "#FFF8EE", "#C55A11"),
        ("实体化数据", 710, 720, "#F8FBFF", "#2F75B5"),
    ]
    for text, x, y, fill, outline in items:
        draw_box(d, (x, y, x + 300, y + 125), text, fill=fill, outline=outline)
        arrow(d, (x + 150, y + 125 if y < 420 else y), (850, 420 if y < 420 else 560))
    img.save(path)


def make_requirement_flow(path: Path):
    img, d = image_base("学生与教师核心业务流程")
    y1, y2 = 190, 560
    student = ["登录", "学习资源", "AI问答", "测验", "画像诊断", "路径/作业"]
    teacher = ["登录", "班级看板", "布置作业", "AI批改", "AI干预", "跟踪反馈"]
    for i, name in enumerate(student):
        x = 80 + i * 260
        draw_box(d, (x, y1, x + 190, y1 + 90), name, fill="#F2F7F2", outline="#548235", size=23)
        if i:
            arrow(d, (x - 70, y1 + 45), (x, y1 + 45))
    for i, name in enumerate(teacher):
        x = 80 + i * 260
        draw_box(d, (x, y2, x + 190, y2 + 90), name, fill="#EAF3F8", outline="#2F75B5", size=23)
        if i:
            arrow(d, (x - 70, y2 + 45), (x, y2 + 45))
    draw_box(d, (650, 385, 1010, 475), "数据沉淀：测验、画像、作业、日志", fill="#FFF8EE", outline="#C55A11")
    arrow(d, (850, y1 + 90), (830, 385))
    arrow(d, (850, 475), (830, y2))
    img.save(path)


def make_architecture_image(path: Path):
    img, d = image_base("系统总体架构设计图")
    layers = [
        ("用户层", 110, 150, ["学生端", "教师端", "管理端"]),
        ("前端层", 110, 310, ["Vue3/Vite", "Router", "API封装", "ECharts"]),
        ("服务层", 110, 470, ["FastAPI", "认证会话", "业务路由", "静态托管"]),
        ("业务层", 110, 630, ["数字孪生", "测验总结", "作业OJ", "教师看板", "AI干预", "行业情报"]),
        ("数据/外部层", 110, 810, ["SQLite/MySQL", "Chroma", "LLM", "go-judge", "JobSpy"]),
    ]
    for title, x, y, nodes in layers:
        d.text((x, y + 30), title, font=get_font(28, True), fill="#17365D")
        x0 = x + 190
        for i, node in enumerate(nodes):
            draw_box(d, (x0 + i * 215, y, x0 + i * 215 + 180, y + 95), node, size=22)
        if y > 150:
            arrow(d, (850, y - 65), (850, y))
    img.save(path)


def make_module_image(path: Path):
    img, d = image_base("后端模块关系图")
    draw_box(d, (650, 150, 1050, 270), "backend/app.py\nFastAPI统一入口", fill="#EAF3F8")
    modules = [
        ("DigitalTwinModule\n学生/教师数字孪生", 70, 390),
        ("HomeworkModule\n作业与判题", 420, 390),
        ("DashboardModule\n教师看板", 770, 390),
        ("IndustryIntelligenceModule\n行业情报", 1120, 390),
        ("Agent/Quiz/Summary/Plan\nAI学习能力", 250, 650),
        ("TeachingInteraction/Research\n互动与教研", 650, 650),
        ("TeacherInterventionModule\nAI干预任务包", 1050, 650),
    ]
    for text, x, y in modules:
        draw_box(d, (x, y, x + 300, y + 115), text, fill="#FFF8EE", outline="#C55A11", size=22)
        arrow(d, (850, 270), (x + 150, y))
    draw_box(d, (560, 830, 1140, 925), "DatabaseModule + tools\n统一存储、RAG、LLM日志、OCR、会话", fill="#F7F7F7", outline="#7F7F7F", size=24)
    for _, x, y in modules:
        arrow(d, (x + 150, y + 115), (850, 830))
    img.save(path)


def make_database_image(path: Path):
    img, d = image_base("数据库核心实体关系图")
    entities = [
        ("users\n用户", 80, 160), ("sessions\n会话", 80, 390), ("teacher_student_links\n师生关系", 80, 620),
        ("courses\n课程", 465, 160), ("course_nodes\n节点", 465, 390), ("resources\n资源", 465, 620),
        ("quiz_attempts\n测验", 850, 160), ("twin_profiles\n画像主表", 850, 390), ("twin_profile_nodes\n画像节点", 850, 620),
        ("learning_plans\n计划/路径", 1235, 160), ("llm_logs\nLLM日志", 1235, 390), ("homework_*\n作业提交", 1235, 620),
    ]
    for text, x, y in entities:
        draw_box(d, (x, y, x + 285, y + 105), text, fill="#F8FBFF", size=22)
    for s, e in [
        ((365, 212), (465, 212)), ((365, 442), (850, 442)), ((365, 672), (850, 672)),
        ((607, 265), (607, 390)), ((607, 495), (607, 620)),
        ((750, 442), (850, 672)), ((992, 265), (992, 390)), ((992, 495), (992, 620)),
        ((1135, 212), (1235, 212)), ((1135, 442), (1235, 442)), ((1135, 672), (1235, 672)),
    ]:
        arrow(d, s, e)
    img.save(path)


def make_ui_image(path: Path):
    img, d = image_base("三端用户界面结构图")
    draw_box(d, (90, 145, 1610, 235), "统一登录页：账号密码登录、角色识别、当前用户查询、按角色路由跳转", fill="#EAF3F8")
    cols = [
        ("学生端", 90, ["首页", "学习中心", "AI问答/总结", "在线测验", "学习诊断", "作业详情", "行业情报", "个人设置"]),
        ("教师端", 630, ["教师看板", "学生详情", "知识点热力图", "作业中心", "AI干预", "教学互动", "教研协同", "教师画像"]),
        ("管理端", 1170, ["管理看板", "学生列表", "教师列表", "LLM日志", "运行状态"]),
    ]
    for title, x, nodes in cols:
        draw_box(d, (x, 320, x + 430, 410), title, fill="#FFF8EE", outline="#C55A11")
        y = 450
        for n in nodes:
            draw_box(d, (x + 35, y, x + 395, y + 52), n, fill="#FFFFFF", outline="#A6A6A6", size=21)
            y += 62
        arrow(d, (850, 235), (x + 215, 320))
    img.save(path)


def make_deploy_image(path: Path):
    img, d = image_base("部署与运行环境图")
    draw_box(d, (110, 160, 430, 280), "浏览器\n学生/教师/管理员")
    draw_box(d, (640, 160, 980, 280), "前端开发\nVite :5173")
    draw_box(d, (1190, 160, 1510, 280), "后端服务\nFastAPI :8000")
    draw_box(d, (140, 520, 410, 640), "SQLite/MySQL\n业务数据")
    draw_box(d, (520, 520, 790, 640), "Chroma\n向量检索")
    draw_box(d, (900, 520, 1170, 640), "go-judge\n代码判题")
    draw_box(d, (1280, 520, 1550, 640), "LLM/JobSpy\n外部服务")
    arrow(d, (430, 220), (640, 220))
    arrow(d, (980, 220), (1190, 220))
    for x in [275, 655, 1035, 1415]:
        arrow(d, (1350, 280), (x, 520))
    draw_box(d, (520, 780, 1180, 890), "发布物：frontend-vue/dist、release/init_seed.sql、config/app_runtime.json、.env", fill="#F7F7F7", outline="#7F7F7F")
    img.save(path)


def make_test_image(path: Path):
    img, d = image_base("系统测试与验收分层图")
    draw_box(d, (550, 150, 1150, 250), "验收测试\n学生学习闭环 / 教师教学闭环 / 管理查看闭环")
    draw_box(d, (420, 340, 1280, 440), "集成测试\nVue 路由 + FastAPI 接口 + DatabaseStore + 外部服务")
    draw_box(d, (270, 530, 1430, 630), "模块测试\n数字孪生 / 作业 / 测验 / RAG / 行业情报 / 干预 / 教学互动")
    draw_box(d, (140, 720, 1560, 820), "单元测试\n计算规则 / Service / Repository / API 参数 / 数据迁移脚本")
    arrow(d, (850, 720), (850, 630))
    arrow(d, (850, 530), (850, 440))
    arrow(d, (850, 340), (850, 250))
    img.save(path)


def setup_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.3)
    sec.right_margin = Cm(2.3)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.bold = True
        style.font.color.rgb = TITLE_COLOR


def shade(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]):
    headers = list(headers)
    rows = [list(row) for row in rows]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, "D9EAF7")
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def bullets(doc: Document, items: Iterable[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def pic(doc: Document, image: Path, width=15.7):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Cm(width))


def cover(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Education 智能教育系统")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle).font.size = Pt(13)
    doc.add_paragraph()
    table(doc, ["项目", "内容"], [
        ("版本", "V2.0"),
        ("日期", "2026-05-13"),
        ("依据", "当前项目代码、接口文档、数据库设计、前端路由与模板文档结构"),
        ("范围", "学生端、教师端、管理端、AI能力、数字孪生、作业、测试、数据与部署"),
    ])
    doc.add_page_break()


def make_requirements_doc(images: dict[str, Path]) -> Path:
    doc = Document()
    setup_doc(doc)
    cover(doc, "产品需求规格说明书", "整体系统详细版")

    doc.add_heading("1 引言", 1)
    doc.add_heading("1.1 编写目的", 2)
    doc.add_paragraph("本文档定义 AI-Education 智能教育系统的整体产品需求，作为后续详细设计、开发实现、测试验收和项目汇报的依据。文档内容以当前项目实现为基础，同时整理可验收的功能边界。")
    doc.add_heading("1.2 项目背景", 2)
    doc.add_paragraph("系统面向大数据课程学习与教学管理场景，围绕课程知识图谱和学习过程数据，构建学生学习闭环、教师教学闭环和管理运维闭环。平台通过 LLM、RAG、数字孪生、作业判题和岗位情报分析提升个性化学习与教学决策效率。")
    doc.add_heading("1.3 业务范围", 2)
    pic(doc, images["scope"])
    table(doc, ["范围", "说明"], [
        ("学生学习", "课程资源、AI 问答、章节总结、在线测验、学习计划、数字孪生诊断、作业、行业情报。"),
        ("教师教学", "班级看板、学生画像、知识点热力图、资源管理、作业发布批改、AI 干预、教学互动、教研协同。"),
        ("管理运维", "用户、教师、学生、LLM 调用日志和运行数据查看。"),
        ("智能能力", "LLM 问答、RAG 检索、OCR、测验生成、AI 批改、路径推荐、岗位技能分析。"),
        ("数据能力", "SQLite/MySQL 实体表、课程图谱、画像历史、学习计划、日志和发布种子。"),
    ])

    doc.add_heading("2 用户角色与业务流程", 1)
    doc.add_heading("2.1 用户角色", 2)
    table(doc, ["角色", "目标", "主要功能"], [
        ("学生", "高效完成课程学习并获得个性化反馈", "学习资源、AI 问答、测验、总结、计划、作业、学习诊断、行业情报。"),
        ("教师", "掌握班级学情并进行精准教学干预", "教师看板、学生详情、热力图、作业中心、AI 批改、干预包、互动、教研、教师画像。"),
        ("管理员", "查看平台基础运行数据", "用户列表、教师列表、学生列表、LLM 日志。"),
        ("系统服务", "自动采集、计算、分析和记录数据", "RAG、LLM、画像采集、OJ 判题、行业抓取、日志记录。"),
    ])
    doc.add_heading("2.2 核心业务流程", 2)
    pic(doc, images["requirement_flow"])

    doc.add_heading("3 功能需求", 1)
    doc.add_paragraph("系统功能按学生端、教师端、管理端和系统服务四类组织。前端页面围绕三类用户建立独立导航，后端接口按业务域封装。")
    pic(doc, images["ui"])
    doc.add_heading("3.1 认证与账户", 2)
    table(doc, ["编号", "需求", "优先级", "验收标准"], [
        ("REQ-AUTH-01", "支持学生、教师、管理员统一登录。", "高", "登录成功后设置 session_id Cookie，并返回当前用户信息。"),
        ("REQ-AUTH-02", "根据 user_type 自动进入对应端。", "高", "student 到学生首页，teacher 到教师看板，admin 到管理看板。"),
        ("REQ-AUTH-03", "支持退出登录和会话清理。", "高", "退出后再次访问受保护页面跳转登录页。"),
        ("REQ-AUTH-04", "支持学生资料维护和密码修改。", "中", "资料更新后 current-user 返回新信息；旧密码错误时拒绝修改。"),
    ])
    doc.add_heading("3.2 学生端功能", 2)
    table(doc, ["编号", "需求", "输入/触发", "输出/结果"], [
        ("REQ-STU-01", "查看课程知识图谱和学习节点。", "进入学习中心", "返回课程结构、节点列表和完成状态。"),
        ("REQ-STU-02", "查看节点资源与 PDF。", "选择知识节点或 PDF", "展示资源列表并可读取 PDF 文件。"),
        ("REQ-STU-03", "AI 课程问答。", "输入问题和会话历史", "结合 RAG 上下文返回回答并记录 LLM 日志。"),
        ("REQ-STU-04", "章节总结。", "输入主题", "生成复习总结。"),
        ("REQ-STU-05", "在线测验。", "选择知识点并答题", "生成题目、判断答案、完成后写入 quiz_attempts。"),
        ("REQ-STU-06", "学习计划。", "输入目标、优先级、截止天数", "生成并保存学习计划。"),
        ("REQ-STU-07", "学习诊断。", "进入诊断页或点击刷新", "展示雷达图、趋势、风险和薄弱点。"),
        ("REQ-STU-08", "作业提交。", "打开作业详情并提交答案/代码", "保存提交，代码题返回判题结果。"),
        ("REQ-STU-09", "行业情报分析。", "输入关键词、地区、数量和阈值", "展示岗位、技能排行、经验学历分布和热力图。"),
    ])
    doc.add_heading("3.3 教师端功能", 2)
    table(doc, ["编号", "需求", "说明", "验收标准"], [
        ("REQ-TEA-01", "班级概览", "查看班级掌握度、风险和学生数量。", "教师进入看板可看到班级摘要。"),
        ("REQ-TEA-02", "学生详情与趋势", "查看单个学生画像和趋势。", "学生详情接口返回画像和趋势数据。"),
        ("REQ-TEA-03", "知识点热力图", "查看节点平均掌握度和学习人数。", "热力图按掌握度排序展示。"),
        ("REQ-TEA-04", "作业管理", "创建、编辑、发布、关闭、重开作业。", "状态变更后学生侧可见性同步变化。"),
        ("REQ-TEA-05", "AI 批改与教师终审", "对提交记录生成 AI 建议并录入最终成绩。", "提交记录包含 ai_score 与 teacher_score。"),
        ("REQ-TEA-06", "AI 干预任务包", "诊断学生、生成草稿、编辑推送、评分反馈。", "学生可接收并提交干预任务。"),
        ("REQ-TEA-07", "教学互动", "公告、话题、帖子、学生提问和教师回复。", "互动统计包含问题数、回复数和平均响应时间。"),
        ("REQ-TEA-08", "教研协同", "记录教研活动、资源链接和课程上下文。", "教研记录可增删改查并同步教师事件。"),
        ("REQ-TEA-09", "教师数字孪生", "汇总教学、互动、教研、批改、干预行为。", "教师画像接口返回维度评分和 AI 建议。"),
    ])
    doc.add_heading("3.4 管理端功能", 2)
    table(doc, ["编号", "需求", "验收标准"], [
        ("REQ-ADM-01", "查看学生列表。", "管理员页面能展示学生账号、姓名、教师等信息。"),
        ("REQ-ADM-02", "查看教师列表。", "管理员页面能展示教师基础信息。"),
        ("REQ-ADM-03", "查看 LLM 调用日志。", "日志包含用户、模块、模型、时间和 payload 摘要。"),
    ])
    doc.add_heading("3.5 数据与智能能力需求", 2)
    table(doc, ["编号", "需求", "说明"], [
        ("REQ-DATA-01", "课程实体化", "课程、节点、资源、测验优先写入实体表。"),
        ("REQ-DATA-02", "学生数字孪生", "采集进度、测验、交互、时长，生成画像和趋势。"),
        ("REQ-DATA-03", "RAG 检索", "资源上传后可进入向量库，问答时检索上下文。"),
        ("REQ-DATA-04", "LLM 日志", "所有主要模型调用记录到 llm_logs。"),
        ("REQ-DATA-05", "OJ 判题", "代码题通过 go-judge 沙箱隔离运行。"),
        ("REQ-DATA-06", "行业情报", "支持岗位抓取、严格相关性过滤和 LLM 技能抽取。"),
    ])

    doc.add_heading("4 非功能需求", 1)
    table(doc, ["类别", "需求"], [
        ("性能", "课程图谱读取使用缓存；大响应启用 GZip；Agent 懒加载；行业情报支持并行采集。"),
        ("可用性", "接口错误应返回明确状态码；前端应提供加载、错误和空数据状态。"),
        ("安全", "受保护接口依赖 session_id；教师资源校验创建者或关联关系；用户代码必须沙箱执行。"),
        ("可维护性", "业务模块按目录拆分；前端 API 和类型集中封装；数据库访问统一通过 DatabaseModule。"),
        ("可观测性", "记录后端日志、LLM 日志、数据迁移报告和接口自测结果。"),
        ("兼容性", "支持 Windows + PowerShell 本地开发；默认 SQLite，可扩展 MySQL。"),
    ])

    doc.add_heading("5 数据需求", 1)
    pic(doc, images["database"])
    table(doc, ["数据对象", "主要字段/内容", "用途"], [
        ("用户", "user_id、login_id、username、user_type、payload_json", "登录、角色和权限判断。"),
        ("课程", "course_id、course_name、payload_json", "知识图谱和学习节点。"),
        ("资源", "course_id、node_id、resource_path、resource_type", "学习资源展示和 RAG 入库。"),
        ("测验", "username、course_id、node_id、score、total、passed", "学习评价和画像同步。"),
        ("学生画像", "overall_mastery、knowledge_nodes、trend", "学习诊断和路径推荐。"),
        ("作业", "assignment、submission、answers、grade", "作业发布、提交和批改。"),
        ("日志", "timestamp、username、module、model、payload_json", "模型调用审计和互动指标。"),
    ])

    doc.add_heading("6 约束、依赖与验收", 1)
    doc.add_heading("6.1 外部依赖", 2)
    bullets(doc, [
        "LLM 服务：通过 .env 配置 model_name、base_url、api_key、embedding_model。",
        "RAG 向量库：依赖 data/chroma_db 与资源入库流程。",
        "OJ 沙箱：依赖 go-judge-sandbox Docker 服务。",
        "行业情报：依赖 JobSpy 和外部招聘网站可用性。",
    ])
    doc.add_heading("6.2 验收标准", 2)
    table(doc, ["闭环", "验收动作", "通过标准"], [
        ("学生学习闭环", "学习资源 -> AI 问答 -> 测验 -> 画像 -> 路径", "各步骤可完成，画像数据更新。"),
        ("教师教学闭环", "看板 -> 作业 -> 学生提交 -> AI 批改 -> 终审", "作业状态和成绩正确。"),
        ("AI 干预闭环", "诊断 -> 生成干预包 -> 推送 -> 学生处理 -> 教师评分", "双方页面状态一致。"),
        ("行业情报闭环", "发起分析 -> 查询状态 -> 展示图表", "结果包含岗位和技能统计。"),
        ("管理闭环", "登录管理端 -> 查看用户和日志", "数据可加载且权限正确。"),
    ])
    path = OUT_DIR / "AI-Education整体系统产品需求规格说明书.docx"
    doc.save(path)
    return path


def make_design_doc(images: dict[str, Path]) -> Path:
    doc = Document()
    setup_doc(doc)
    cover(doc, "系统设计说明书", "整体系统详细版")

    doc.add_heading("1 设计概述", 1)
    doc.add_paragraph("系统采用前后端分离和领域模块化设计。前端承担路由、交互和图表展示；后端承担统一认证、业务编排、数据持久化、AI 能力调用与静态资源托管；数据层通过 Store 抽象支持 SQLite 和 MySQL。")
    pic(doc, images["architecture"])
    doc.add_heading("1.1 设计目标", 2)
    bullets(doc, [
        "保持学生端、教师端、管理端角色边界清晰。",
        "通过实体化数据库降低 JSON 文件并存阶段的数据不一致风险。",
        "将 LLM、RAG、OJ、岗位抓取等外部能力封装为可替换服务。",
        "通过数字孪生沉淀学习与教学过程数据，支撑诊断和干预。",
        "保证本地开发可运行、团队协作可复现、发布数据可同步。",
    ])

    doc.add_heading("2 前端设计", 1)
    pic(doc, images["ui"])
    doc.add_heading("2.1 路由结构", 2)
    table(doc, ["端", "路由", "页面职责"], [
        ("学生端", "/student/home", "学习首页和概览。"),
        ("学生端", "/student/course-content", "课程内容、资源、AI 问答和 PDF 阅读。"),
        ("学生端", "/student/student-twin", "学习诊断、雷达图、趋势、风险和薄弱点。"),
        ("学生端", "/student/quiz", "在线测验。"),
        ("学生端", "/student/homework", "作业列表和提交。"),
        ("学生端", "/student/industry-intelligence", "行业岗位和技能分析。"),
        ("教师端", "/teacher/dashboard", "班级看板和教师画像入口。"),
        ("教师端", "/teacher/homework", "作业发布、提交查看和批改。"),
        ("教师端", "/teacher/intervention", "AI 干预任务包。"),
        ("教师端", "/teacher/interaction", "公告、讨论和互动统计。"),
        ("教师端", "/teacher/research", "教研协同记录。"),
        ("管理端", "/admin/dashboard", "用户、教师、学生、LLM 日志。"),
    ])
    doc.add_heading("2.2 前端模块约定", 2)
    table(doc, ["目录/文件", "职责"], [
        ("src/router/index.ts", "路由注册、角色跳转和登录守卫。"),
        ("src/layouts", "StudentLayout、TeacherLayout、AdminLayout 三端布局。"),
        ("src/api", "按业务域封装请求，如 homework、student、teacher、industry、intervention。"),
        ("src/types", "前端类型定义。"),
        ("src/views/student", "学生端页面。"),
        ("src/views/teacher", "教师端页面。"),
        ("src/views/admin", "管理端页面。"),
    ])

    doc.add_heading("3 后端架构设计", 1)
    pic(doc, images["module"])
    doc.add_heading("3.1 FastAPI 入口", 2)
    bullets(doc, [
        "backend/app.py 创建 FastAPI 应用并注册中间件。",
        "include_router 注册数字孪生、看板、作业、行业情报、教学互动、教研和干预路由。",
        "挂载 /static、/data、/assets，并提供前端 SPA 回退。",
        "启动时预热 RAG，并创建数字孪生定时采集任务。",
        "Agent 采用懒加载，避免启动阶段耗时过长。",
    ])
    doc.add_heading("3.2 业务模块设计", 2)
    table(doc, ["模块", "核心类/文件", "说明"], [
        ("认证与会话", "UserManager、SessionManager", "登录、当前用户、Cookie 会话和角色信息。"),
        ("课程与资源", "backend/app.py、tools/rag_service", "课程图谱、学习节点、资源上传删除、PDF 读取和 RAG 入库。"),
        ("学生数字孪生", "DataCollector、ScoreCalculator、StudentTwinService", "采集学习行为并输出画像摘要。"),
        ("教师数字孪生", "TeacherTwinService、teacher_event_repository", "汇总教师互动、教研、批改和干预行为。"),
        ("教师看板", "DashboardModule", "班级概览、学生趋势、节点排名和教师画像建议。"),
        ("作业中心", "HomeworkModule", "作业、提交、AI 批改、终审、OJ 判题。"),
        ("AI 干预", "TeacherInterventionModule", "诊断、草稿、推送、学生处理和教师评分。"),
        ("教学互动", "TeachingInteractionModule", "公告、讨论、帖子、学生提问和教师回复。"),
        ("教研协同", "TeachingResearchModule", "教研记录和教师事件沉淀。"),
        ("行业情报", "IndustryIntelligenceModule", "岗位抓取、相关性过滤、技能分析和图表数据。"),
    ])

    doc.add_heading("4 数据库设计", 1)
    pic(doc, images["database"])
    doc.add_heading("4.1 数据表分组", 2)
    table(doc, ["分组", "表", "说明"], [
        ("用户权限", "users、sessions、teacher_student_links、user_states", "用户、角色、会话、师生关系和扩展状态。"),
        ("课程资源", "courses、course_nodes、resources", "课程图谱、节点、PDF/视频/资料资源。"),
        ("学习评价", "quiz_attempts、learning_plans、llm_logs", "测验、计划/路径、模型调用日志。"),
        ("数字孪生", "twin_profiles、twin_profile_nodes、twin_history", "学生画像主表、节点明细、趋势历史。"),
        ("作业", "homework_assignments、homework_submissions", "作业题目、学生答案、AI 批改和教师终审。"),
    ])
    doc.add_heading("4.2 实体化读写策略", 2)
    table(doc, ["对象", "读策略", "写策略"], [
        ("课程图谱", "优先 courses.payload_json", "sync_course_from_graph 同步实体表。"),
        ("学习节点", "优先 course_nodes", "课程同步时拆分节点。"),
        ("资源", "优先 resources", "上传/删除时实体表主写，并同步课程 payload。"),
        ("测验", "quiz_attempts 查询历史", "测验完成写实体表并同步画像。"),
        ("画像", "twin_profiles + twin_profile_nodes", "采集后统一 recalculation 并保存。"),
    ])

    doc.add_heading("5 接口设计", 1)
    table(doc, ["接口域", "代表接口", "说明"], [
        ("认证", "POST /api/auth/login、GET /api/current-user、POST /api/logout", "统一登录和会话。"),
        ("课程", "GET /api/knowledge-graph、GET /api/learning-nodes、POST /api/node/resources", "课程图谱、节点和资源。"),
        ("AI学习", "POST /api/chat、POST /api/summary、POST /api/quiz/start、POST /api/quiz/complete", "问答、总结、测验。"),
        ("数字孪生", "GET /api/digital-twin/student-profile/{username}、POST /api/digital-twin/student-course-profile", "画像摘要和课程画像。"),
        ("作业", "GET/POST /api/homework/assignments、POST /api/homework/submissions/{id}/ai-grade", "作业管理和批改。"),
        ("教师看板", "GET /api/dashboard/class-overview、GET /api/heatmap", "班级学情和热力图。"),
        ("行业情报", "POST /api/industry-intelligence/analyze、GET /api/industry-intelligence/tasks/{task_id}", "任务型岗位分析。"),
    ])

    doc.add_heading("6 核心流程设计", 1)
    doc.add_heading("6.1 学生学习闭环", 2)
    table(doc, ["步骤", "前端动作", "后端处理", "数据变化"], [
        ("1", "打开学习中心", "读取课程图谱和节点资源", "courses、course_nodes、resources"),
        ("2", "提问或总结", "RAG 检索 + LLM 生成", "llm_logs 增加记录"),
        ("3", "完成测验", "记录测验并判断通过", "quiz_attempts 写入"),
        ("4", "刷新诊断", "采集进度/日志/会话并计算画像", "twin_profiles、twin_profile_nodes、twin_history 更新"),
        ("5", "生成路径", "弱项检测 + 资源推荐 + LLM 排序", "learning_plans 写入 path 分类"),
    ])
    doc.add_heading("6.2 教师作业闭环", 2)
    table(doc, ["步骤", "处理说明"], [
        ("创建作业", "教师录入作业、题目、课程节点、评分标准，可保存草稿或发布。"),
        ("学生提交", "学生提交主观题、客观题或代码答案，代码题可走 OJ 沙箱。"),
        ("AI 批改", "教师触发 AI 批改，系统生成建议分、反馈和依据。"),
        ("教师终审", "教师录入最终成绩和评语，形成最终评价。"),
        ("画像联动", "作业表现可作为后续学生和教师数字孪生扩展指标。"),
    ])

    doc.add_heading("7 部署设计", 1)
    pic(doc, images["deploy"])
    table(doc, ["组件", "配置/命令", "说明"], [
        ("后端", "python main.py", "默认 8000 端口。"),
        ("前端开发", "cd frontend-vue && npm run dev", "默认 5173 端口，经 Vite 代理访问后端。"),
        ("前端生产", "npm run build", "后端托管 frontend-vue/dist。"),
        ("OJ 沙箱", "docker compose up -d --build", "go-judge-sandbox 提供 /run。"),
        ("数据库", "DB_TYPE、DB_PATH 或 MySQL 环境变量", "默认 data/app.db。"),
        ("LLM", "model_name、base_url、api_key、embedding_model", "OpenAI 兼容接口。"),
    ])

    doc.add_heading("8 安全、性能与可维护性设计", 1)
    table(doc, ["方面", "设计"], [
        ("安全", "HttpOnly Cookie；接口角色校验；教师资源所有者校验；用户代码沙箱运行；API Key 不入库。"),
        ("性能", "GZip 响应压缩；课程缓存；Agent 懒加载；RAG 启动预热；行业情报并行抓取。"),
        ("可维护", "业务模块目录化；API 封装与类型定义分离；数据库 Store 抽象；文档同步维护。"),
        ("可观测", "后端日志、LLM 日志、数据迁移报告、接口自测脚本。"),
    ])

    doc.add_heading("9 测试与验收设计", 1)
    pic(doc, images["test"])
    table(doc, ["测试层级", "测试对象", "建议命令/方式"], [
        ("单元测试", "ScoreCalculator、Service、Repository、路径推荐", "pytest tests"),
        ("接口测试", "登录、课程、测验、画像、作业、看板", "python scripts\\backend_api_selftest.py --base-url http://127.0.0.1:8000"),
        ("前端构建", "Vue 类型和构建", "cd frontend-vue && npm run build"),
        ("手工验收", "学生、教师、管理核心闭环", "按业务闭环逐项操作。"),
        ("沙箱测试", "代码题判题", "scripts\\test_sandbox_languages.py 或 OJ smoke assignment。"),
    ])

    doc.add_heading("10 演进建议", 1)
    bullets(doc, [
        "将历史 JSON 兜底导入迁移为独立命令，运行时彻底实体表优先。",
        "完善多课程与班级绑定，支持不同课程画像和资源隔离。",
        "补充前端 Playwright 自动化测试，覆盖学生诊断、作业提交和教师批改。",
        "增加统一任务表，管理行业情报、资源入库、画像重算等后台长任务。",
        "统一修复旧文件编码问题，提升后续维护效率。",
    ])
    path = OUT_DIR / "AI-Education整体系统设计说明书.docx"
    doc.save(path)
    return path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "scope": IMG_DIR / "scope.png",
        "requirement_flow": IMG_DIR / "requirement_flow.png",
        "architecture": IMG_DIR / "architecture.png",
        "module": IMG_DIR / "module.png",
        "database": IMG_DIR / "database.png",
        "ui": IMG_DIR / "ui.png",
        "deploy": IMG_DIR / "deploy.png",
        "test": IMG_DIR / "test.png",
    }
    make_scope_image(images["scope"])
    make_requirement_flow(images["requirement_flow"])
    make_architecture_image(images["architecture"])
    make_module_image(images["module"])
    make_database_image(images["database"])
    make_ui_image(images["ui"])
    make_deploy_image(images["deploy"])
    make_test_image(images["test"])
    req = make_requirements_doc(images)
    design = make_design_doc(images)
    print(req)
    print(design)


if __name__ == "__main__":
    main()
