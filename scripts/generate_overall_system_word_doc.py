from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "整体系统Word文档"
IMG_DIR = OUT_DIR / "images"
TITLE_COLOR = RGBColor(31, 78, 121)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines, current = [], ""
    for ch in text:
        trial = current + ch
        bbox = draw.textbbox((0, 0), trial, font=fnt)
        if bbox[2] - bbox[0] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def canvas(title: str, w=1600, h=950):
    img = Image.new("RGB", (w, h), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.text((56, 34), title, font=font(40, True), fill="#17365D")
    d.line((56, 92, w - 56, 92), fill="#9DB7D5", width=3)
    return img, d


def box(draw, xy, text, fill="#F7FBFF", outline="#2F75B5", text_color="#1F2937", size=26):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    fnt = font(size, True)
    lines = wrap_text(draw, text, fnt, x2 - x1 - 28)
    line_h = int(size * 1.28)
    total_h = len(lines) * line_h
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=fnt, fill=text_color)
        y += line_h


def arrow(draw, start, end, color="#6B7280", width=4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - sign * 16, y2 - 9), (x2 - sign * 16, y2 + 9)]
    else:
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - sign * 16), (x2 + 9, y2 - sign * 16)]
    draw.polygon(pts, fill=color)


def make_architecture(path: Path):
    img, d = canvas("AI-Education 整体系统架构图")
    box(d, (80, 150, 390, 270), "学生端\n学习/测验/作业/诊断")
    box(d, (80, 330, 390, 450), "教师端\n看板/作业/干预/教研")
    box(d, (80, 510, 390, 630), "管理端\n用户/日志/运行查看")
    box(d, (585, 230, 930, 380), "Vue 3 前端\nRouter + API 封装\nElement Plus + ECharts", fill="#F2F7F2", outline="#548235")
    box(d, (1120, 230, 1500, 380), "FastAPI 后端\n认证/课程/AI/作业/孪生/行业情报", fill="#EAF3F8")
    modules = [
        ("数字孪生", 190, 735),
        ("AI Agent/RAG", 410, 735),
        ("作业/OJ", 630, 735),
        ("教师看板", 850, 735),
        ("教学互动/教研", 1070, 735),
        ("行业情报", 1290, 735),
    ]
    for name, x, y in modules:
        box(d, (x, y, x + 190, y + 90), name, fill="#FFF8EE", outline="#C55A11", size=24)
        arrow(d, (1310, 380), (x + 95, y))
    box(d, (555, 535, 1010, 650), "DatabaseModule\nSQLite 默认 / MySQL 可选", fill="#F7F7F7", outline="#7F7F7F")
    box(d, (1125, 535, 1485, 650), "外部能力\nLLM / Chroma / go-judge / JobSpy", fill="#FFF2F2", outline="#C00000")
    for y in [210, 390, 570]:
        arrow(d, (390, y), (585, 305))
    arrow(d, (930, 305), (1120, 305))
    arrow(d, (1310, 380), (1310, 535))
    arrow(d, (1120, 335), (1010, 590))
    img.save(path)


def make_business(path: Path):
    img, d = canvas("核心业务闭环图")
    steps = [
        ("课程学习", 90, 190),
        ("AI 问答", 330, 190),
        ("在线测验", 570, 190),
        ("画像更新", 810, 190),
        ("路径推荐", 1050, 190),
        ("作业反馈", 1290, 190),
    ]
    for name, x, y in steps:
        box(d, (x, y, x + 190, y + 95), name)
    for i in range(len(steps) - 1):
        arrow(d, (steps[i][1] + 190, steps[i][2] + 48), (steps[i + 1][1], steps[i + 1][2] + 48))
    box(d, (210, 520, 520, 640), "教师看板\n掌握班级学情", fill="#F2F7F2", outline="#548235")
    box(d, (645, 520, 955, 640), "AI 干预任务包\n精准补救", fill="#FFF8EE", outline="#C55A11")
    box(d, (1080, 520, 1390, 640), "教学互动/教研\n沉淀教师行为", fill="#F2F7F2", outline="#548235")
    arrow(d, (900, 285), (365, 520))
    arrow(d, (900, 285), (800, 520))
    arrow(d, (1385, 285), (1235, 520))
    box(d, (560, 760, 1040, 850), "数据沉淀：课程、资源、测验、画像、作业、日志", fill="#F7F7F7", outline="#7F7F7F")
    arrow(d, (365, 640), (690, 760))
    arrow(d, (800, 640), (800, 760))
    arrow(d, (1235, 640), (910, 760))
    img.save(path)


def make_database(path: Path):
    img, d = canvas("核心数据库实体关系图")
    entities = [
        ("users\n用户主数据", 70, 160, 310, 270),
        ("sessions\n会话状态", 70, 380, 310, 490),
        ("courses\n课程主表", 470, 160, 710, 270),
        ("course_nodes\n课程节点", 470, 380, 710, 490),
        ("resources\n学习资源", 470, 600, 710, 710),
        ("quiz_attempts\n测验记录", 870, 160, 1110, 270),
        ("twin_profiles\n画像主表", 870, 380, 1110, 490),
        ("twin_profile_nodes\n画像节点", 870, 600, 1110, 710),
        ("learning_plans\n学习计划/路径", 1270, 160, 1510, 270),
        ("llm_logs\n模型调用日志", 1270, 380, 1510, 490),
        ("homework_*\n作业与提交", 1270, 600, 1510, 710),
    ]
    for text, x1, y1, x2, y2 in entities:
        box(d, (x1, y1, x2, y2), text, fill="#F8FBFF")
    for s, e in [
        ((310, 215), (470, 215)),
        ((590, 270), (590, 380)),
        ((590, 490), (590, 600)),
        ((710, 435), (870, 655)),
        ((990, 270), (990, 380)),
        ((990, 490), (990, 600)),
        ((1110, 435), (1270, 435)),
        ((1110, 215), (1270, 215)),
        ((1110, 655), (1270, 655)),
        ((310, 435), (870, 435)),
    ]:
        arrow(d, s, e)
    img.save(path)


def make_ui(path: Path):
    img, d = canvas("三端界面结构图")
    box(d, (80, 145, 1480, 235), "统一登录页：账号、密码、角色识别、登录后按 user_type 跳转")
    columns = [
        ("学生端", 80, ["首页", "学习中心", "学习诊断", "测验", "作业", "行业情报", "个人设置"]),
        ("教师端", 580, ["教师看板", "教学互动", "教研协同", "作业中心", "AI干预", "教师画像钻取"]),
        ("管理端", 1080, ["管理看板", "用户列表", "教师/学生", "LLM日志", "运行数据"]),
    ]
    for title, x, items in columns:
        box(d, (x, 330, x + 400, 430), title, fill="#EAF3F8")
        y = 470
        for item in items:
            box(d, (x + 35, y, x + 365, y + 58), item, fill="#FFFFFF", outline="#A6A6A6", size=22)
            y += 70
        arrow(d, (780, 235), (x + 200, 330))
    img.save(path)


def make_deploy(path: Path):
    img, d = canvas("运行部署与外部依赖图")
    box(d, (80, 170, 420, 300), "前端开发服务\nVite :5173")
    box(d, (590, 170, 930, 300), "后端服务\nFastAPI :8000")
    box(d, (1100, 170, 1440, 300), "前端生产包\nfrontend-vue/dist")
    box(d, (80, 520, 360, 650), "SQLite\nrelease/init_seed.sql")
    box(d, (440, 520, 720, 650), "Chroma\n向量库")
    box(d, (800, 520, 1080, 650), "go-judge\nDocker沙箱")
    box(d, (1160, 520, 1440, 650), "LLM/招聘网站\n外部服务")
    arrow(d, (420, 235), (590, 235))
    arrow(d, (930, 235), (1100, 235))
    for x in [220, 580, 940, 1300]:
        arrow(d, (760, 300), (x, 520))
    img.save(path)


def make_test(path: Path):
    img, d = canvas("系统测试分层图")
    box(d, (560, 140, 1040, 240), "验收测试\n学生闭环 / 教师闭环 / 管理闭环")
    box(d, (430, 330, 1170, 430), "集成测试\n前端路由 + FastAPI + Store + 外部服务")
    box(d, (280, 520, 1320, 620), "模块测试\n数字孪生 / 作业 / 测验 / RAG / 行业情报 / 干预")
    box(d, (150, 710, 1450, 810), "单元测试\n计算规则、Repository、Service、API 参数校验、数据迁移脚本")
    arrow(d, (800, 710), (800, 620))
    arrow(d, (800, 520), (800, 430))
    arrow(d, (800, 330), (800, 240))
    img.save(path)


def set_style(doc: Document):
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.bold = True
        style.font.color.rgb = TITLE_COLOR


def shade(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]):
    headers = list(headers)
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        shade(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_pic(doc: Document, path: Path, width=15.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))


def bullets(doc: Document, items: Iterable[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc(images: dict[str, Path]) -> Path:
    doc = Document()
    set_style(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Education 智能教育系统")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("整体系统需求与设计说明书")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = TITLE_COLOR
    doc.add_paragraph()
    meta = add_table(doc, ["项目", "内容"], [
        ("版本", "V1.0"),
        ("日期", "2026-05-13"),
        ("技术栈", "FastAPI + Vue 3 + SQLite/MySQL + LLM + RAG"),
        ("覆盖范围", "学生端、教师端、管理端、数字孪生、作业、测验、行业情报、教学互动、教研协同"),
    ])
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()

    doc.add_heading("1 项目概述", 1)
    doc.add_paragraph("AI-Education 是面向课程学习和教学管理的智能教育平台。系统以课程知识图谱为基础，以学生数字孪生和教师数字孪生为数据中枢，串联学习、测验、作业、AI 问答、个性化路径、教师看板、AI 干预、教学互动、教研协同和行业情报分析。")
    add_pic(doc, images["business"])

    doc.add_heading("2 项目计划与交付范围", 1)
    add_table(doc, ["阶段", "主要任务", "交付物"], [
        ("需求阶段", "明确三端角色、业务闭环和核心数据", "需求规格、功能清单"),
        ("设计阶段", "完成前后端、数据库、接口和模块设计", "体系结构、数据库、界面设计"),
        ("实现阶段", "开发 FastAPI 接口、Vue 页面和业务模块", "可运行系统与模块代码"),
        ("联调阶段", "完成登录、学习、测验、作业、画像、看板联调", "接口自测与页面验收"),
        ("发布阶段", "准备种子数据、部署配置和运行说明", "release 数据、README、测试报告"),
    ])
    doc.add_heading("3 产品需求", 1)
    add_table(doc, ["角色", "核心功能"], [
        ("学生", "课程学习、AI 问答、章节总结、在线测验、学习计划、学习诊断、作业提交、行业情报。"),
        ("教师", "班级看板、学生详情、知识点热力图、资源管理、作业发布批改、AI 干预、互动与教研、教师画像。"),
        ("管理员", "查看用户、教师、学生、LLM 日志和平台运行数据。"),
        ("系统服务", "数据采集、画像计算、RAG 检索、LLM 调用、OJ 判题、岗位抓取。"),
    ])
    doc.add_heading("4 总体架构设计", 1)
    add_pic(doc, images["architecture"])
    add_table(doc, ["层级", "设计说明"], [
        ("前端层", "Vue 3 + Vite，按学生、教师、管理员拆分布局和路由。"),
        ("接口层", "FastAPI 提供统一 REST 接口和 SPA 静态托管。"),
        ("业务层", "按模块拆分数字孪生、作业、看板、行业情报、教学互动、教研和干预。"),
        ("数据层", "DatabaseModule 统一封装 SQLite/MySQL，实体表优先。"),
        ("智能层", "LLM、RAG、OCR、Agent、go-judge、JobSpy 等能力按需接入。"),
    ])
    doc.add_heading("5 用户界面设计", 1)
    add_pic(doc, images["ui"])
    bullets(doc, [
        "学生端强调学习闭环：首页、学习中心、诊断、测验、作业、行业情报、个人设置。",
        "教师端强调管理闭环：看板、互动、教研、作业、AI 干预、教师画像钻取。",
        "管理端保持简洁：聚焦用户、日志和运行数据查看。",
        "前端统一通过 API 封装访问后端，路由守卫根据当前用户角色跳转。",
    ])
    doc.add_heading("6 数据库设计", 1)
    add_pic(doc, images["database"])
    add_table(doc, ["数据域", "核心表"], [
        ("用户与会话", "users、sessions、teacher_student_links、user_states"),
        ("课程与资源", "courses、course_nodes、resources"),
        ("学习与测验", "quiz_attempts、learning_plans、llm_logs"),
        ("数字孪生", "twin_profiles、twin_profile_nodes、twin_history"),
        ("作业", "homework_assignments、homework_submissions"),
        ("教师事件", "教学互动、教研、批改和干预相关事件表/状态数据"),
    ])
    doc.add_heading("7 核心技术实现", 1)
    add_table(doc, ["模块", "核心技术"], [
        ("学生数字孪生", "多源采集、掌握度加权计算、雷达图、风险预警、趋势分析、路径联动。"),
        ("AI 问答与总结", "RAG 检索课程资源，LLM 生成问答、总结和学习计划。"),
        ("作业中心", "作业状态机、AI 辅助批改、教师终审、go-judge 沙箱判题。"),
        ("教师看板", "班级概览、学生趋势、知识点排名、热力图和教师画像。"),
        ("AI 干预", "基于学生画像生成任务包，支持推送、学生处理、教师评分。"),
        ("行业情报", "JobSpy 抓取岗位，相关性过滤，LLM 提取技能和要求，ECharts 展示。"),
    ])
    doc.add_heading("8 部署与运行设计", 1)
    add_pic(doc, images["deploy"])
    add_table(doc, ["对象", "命令/配置"], [
        ("后端启动", "python main.py 或 uvicorn backend.app:app --host 0.0.0.0 --port 8000 --reload"),
        ("前端开发", "cd frontend-vue && npm run dev"),
        ("前端构建", "cd frontend-vue && npm run build"),
        ("OJ 沙箱", "cd go-judge-sandbox && docker compose up -d --build"),
        ("数据库种子", "release/init_seed.sql、release/app_release.db"),
        ("模型配置", ".env 中配置 model_name、base_url、api_key、embedding_model"),
    ])
    doc.add_heading("9 系统测试设计", 1)
    add_pic(doc, images["test"])
    add_table(doc, ["测试类型", "重点"], [
        ("单元测试", "掌握度计算、路径推荐、作业逻辑、Repository 和 Service。"),
        ("接口测试", "登录、课程、测验、画像、作业、看板、行业情报。"),
        ("前端测试", "学生/教师/管理路由、图表渲染、表单提交和错误提示。"),
        ("集成测试", "测验完成后画像同步，作业提交后批改，教师干预闭环。"),
        ("验收测试", "学生学习闭环、教师教学闭环、管理员查看闭环。"),
    ])
    doc.add_heading("10 风险与改进建议", 1)
    add_table(doc, ["风险", "建议"], [
        ("历史 JSON 与实体表并存", "将兜底导入迁移为显式脚本，运行时以实体表为准。"),
        ("多课程逻辑仍在演进", "完善 course_id 与班级/用户绑定关系。"),
        ("外部 LLM 与招聘网站不稳定", "增加降级提示、缓存和任务状态恢复。"),
        ("部分源码/旧文档存在编码显示问题", "统一 UTF-8 编码并清理历史文本。"),
        ("前端自动化测试不足", "补充 Playwright 核心页面截图与交互测试。"),
    ])

    path = OUT_DIR / "AI-Education整体系统需求与设计说明书.docx"
    doc.save(path)
    return path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        "architecture": IMG_DIR / "overall_architecture.png",
        "business": IMG_DIR / "business_loop.png",
        "database": IMG_DIR / "database_er.png",
        "ui": IMG_DIR / "ui_structure.png",
        "deploy": IMG_DIR / "deploy.png",
        "test": IMG_DIR / "test_layers.png",
    }
    make_architecture(images["architecture"])
    make_business(images["business"])
    make_database(images["database"])
    make_ui(images["ui"])
    make_deploy(images["deploy"])
    make_test(images["test"])
    print(build_doc(images))


if __name__ == "__main__":
    main()

