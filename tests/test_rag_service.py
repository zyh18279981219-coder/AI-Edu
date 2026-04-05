import logging
import base64
import tempfile
import os
from langchain_core.documents import Document
from langchain_community.embeddings import FakeEmbeddings
from langchain_community.document_loaders import UnstructuredFileLoader
from docx import Document as DocxDocument
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tools.rag_service import RAGService

import tools.rag_service as rag_service
from dotenv import load_dotenv
import os

load_dotenv()
model_name = os.environ.get("model_name")
base_url = os.environ.get("base_url")
api_key = os.environ.get("api_key")
embedding_model = os.environ.get("embedding_model")


def test_get_rag_service_lazy_singleton(monkeypatch):
    """Ensure ``get_rag_service`` lazily creates a singleton instance."""
    with tempfile.TemporaryDirectory() as tmpdir:

        class TestService(rag_service.RAGService):
            def __init__(self, *args, **kwargs):

                kwargs["embeddings"] = FakeEmbeddings(size=32)
                kwargs["persist_directory"] = tmpdir
                kwargs["use_multiquery"] = False

                super().__init__(*args, **kwargs)

        monkeypatch.setattr(rag_service, "_instance", None)
        monkeypatch.setattr(rag_service, "RAGService", TestService)

        first = rag_service.get_rag_service()
        second = rag_service.get_rag_service()
        assert first is second


def test_ingest_and_retrieve(tmp_path):

    service = RAGService(
        embeddings=FakeEmbeddings(size=32),
        persist_directory=str(tmp_path),
        use_multiquery=False,
    )
    doc = Document(page_content="Cats are great pets")
    service.ingest_paths([doc])
    retriever = service.get_retriever()
    docs = retriever.get_relevant_documents("cats")
    assert any("Cats are great pets" in d.page_content for d in docs)


def test_ingest_paths_lookup_error(monkeypatch, tmp_path, caplog):
    def bad_load(self):
        raise LookupError("punkt not found")

    monkeypatch.setattr(UnstructuredFileLoader, "load", bad_load)
    file_path = tmp_path / "f.txt"
    file_path.write_text("test")

    service = RAGService(
        embeddings=FakeEmbeddings(size=32),
        persist_directory=str(tmp_path / "db"),
        use_multiquery=False,
    )

    caplog.set_level(logging.ERROR)
    msg = service.ingest_paths([str(file_path)])
    assert "nltk.download('punkt')" in msg
    assert "Failed to load" in caplog.text


def test_ingest_paths_import_error(monkeypatch, tmp_path, caplog):
    def bad_load(self):
        raise ImportError("unstructured dependency missing")

    monkeypatch.setattr(UnstructuredFileLoader, "load", bad_load)
    file_path = tmp_path / "f.txt"
    file_path.write_text("test")

    service = RAGService(
        embeddings=FakeEmbeddings(size=32),
        persist_directory=str(tmp_path / "db2"),
        use_multiquery=False,
    )

    caplog.set_level(logging.ERROR)
    msg = service.ingest_paths([str(file_path)])
    assert "unstructured[pdf]" in msg
    assert "Failed to load" in caplog.text


def test_ingest_pdf(tmp_path):
    pdf_b64 = (
        "JVBERi0xLjUKMSAwIG9iaiA8PCAvVHlwZSAvQ2F0YWxvZyAvUGFnZXMgMiAwIFIgPj4gZW5k"
        "b2JqCjIgMCBvYmogPDwgL1R5cGUgL1BhZ2VzIC9LaWRzIFszIDAgUl0gL0NvdW50IDEgPj4g"
        "ZW5kb2JqCjMgMCBvYmogPDwgL1R5cGUgL1BhZ2UgL1BhcmVudCAyIDAgUiAvTWVkaWFCb3gg"
        "WzAgMCAyMDAgMjAwXSAvQ29udGVudHMgNCAwIFIgL1Jlc291cmNlcyA8PCAvRm9udCA8PCAv"
        "RjEgNSAwIFIgPj4gPj4gPj4gZW5kb2JqCjQgMCBvYmogPDwgL0xlbmd0aCA0NCA+PiBzdHJl"
        "YW0KQlQgL0YxIDEyIFRmIDcyIDEwMCBUZCAoQ2F0cyBhcmUgZ3JlYXQgcGV0cykgVGogRVQK"
        "ZW5kc3RyZWFtIGVuZG9iago1IDAgb2JqIDw8IC9UeXBlIC9Gb250IC9TdWJ0eXBlIC9UeXBl"
        "MSAvQmFzZUZvbnQgL0hlbHZldGljYSA+PiBlbmRvYmoKeHJlZgowIDYKMDAwMDAwMDAwMCA2"
        "NTUzNSBmIAowMDAwMDAwMDEwIDAwMDAwIG4gCjAwMDAwMDAwNTYgMDAwMDAgbiAKMDAwMDAw"
        "MDExNCAwMDAwMCBuIAowMDAwMDAwMjQ1IDAwMDAwIG4gCjAwMDAwMDAzMzQgMDAwMDAgbiAK"
        "dHJhaWxlciA8PCAvUm9vdCAxIDAgUiAvU2l6ZSA2ID4+CnN0YXJ0eHJlZgo0MTYKJSVFT0YK"
    )
    pdf_path = tmp_path / "cats.pdf"
    pdf_path.write_bytes(base64.b64decode(pdf_b64))

    service = RAGService(
        embeddings=FakeEmbeddings(size=32),
        persist_directory=str(tmp_path / "db3"),
        use_multiquery=False,
    )
    err = service.ingest_paths([str(pdf_path)])
    assert err is None
    docs = service.get_retriever().get_relevant_documents("cats")
    assert any("Cats are great pets" in d.page_content for d in docs)


def test_ingest_docx(tmp_path):
    docx_path = tmp_path / "cats.docx"
    doc = DocxDocument()
    doc.add_paragraph("Cats are playful animals")
    doc.save(docx_path)

    service = RAGService(
        embeddings=FakeEmbeddings(size=32),
        persist_directory=str(tmp_path / "db4"),
        use_multiquery=False,
    )
    err = service.ingest_paths([str(docx_path)])
    assert err is None
    docs = service.get_retriever().get_relevant_documents("playful")
    assert any("Cats are playful animals" in d.page_content for d in docs)
